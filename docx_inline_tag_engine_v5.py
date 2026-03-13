"""
docx_inline_tag_engine.py  v5
==============================
模板 + 数据分离架构

v5 新增：目录（TOC）支持
  ① JSON 里加 {"type": "toc"} 块，引擎在该位置生成目录
  ② 双轨渲染模式：
       auto  — 插入 Word TOC 域（打开时自动刷新，格式由 TOC 1/2/3 样式决定）
       manual— 用提取的模板样式手动排版（页码由 JSON 提供，适合精确控制）
  ③ toc_exclude：在 section / heading 块上设为 true，该标题不进入目录
  ④ 模板 [[目录]] 块：放三行示例条目（1/2/3级格式），引擎提取其 pPr/rPr

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

JSON 完整结构（v5）
-------------------
{
  "page_footer_spec": "...",           // 自然语言页脚描述，LLM 解析
  // 或者直接:
  "page_footer_config": [...],

  "toc_mode": "auto",                  // "auto"（默认）或 "manual"
  // manual 模式下需要提供每个标题的页码：
  // "toc_entries": [
  //   {"title": "摘  要",   "level": 1, "page": "i"},
  //   {"title": "第一章 引言", "level": 1, "page": 1},
  //   ...
  // ]

  "content": [
    // ── 前置章节（toc_exclude 控制是否进目录）──────────────────
    {"type": "section", "section_type": "abstract",
     "toc_exclude": true,              // 不进目录（默认 false）
     "value": "本文研究了…"},

    // ── 目录占位（放在想要目录出现的位置）─────────────────────
    {"type": "toc",
     "title": "目  录",                // 可选，目录章节标题
     "toc_title_exclude": true},       // 目录标题本身是否排除出目录（默认 true）

    // ── 正文 ────────────────────────────────────────────────────
    {"type": "heading1", "value": "第一章 引言"},
    {"type": "heading2", "value": "1.1 研究背景"},
    {"type": "body",     "value": "正文…"},
    {"type": "table",    "caption": "表1", "data": [[...]]},
    {"type": "formula",  "label": "式(1)", "latex": "..."},
    {"type": "image",    "path": "fig.png", "caption": "图1"},

    // ── 后置章节 ────────────────────────────────────────────────
    {"type": "section", "section_type": "conclusion",
     "toc_exclude": false,             // 进目录
     "value": "综上所述…"},

    {"type": "section", "section_type": "acknowledgement",
     "toc_exclude": true,
     "value": "感谢…"},

    {"type": "section", "section_type": "publications",
     "toc_exclude": false,
     "value": "1. …"}
  ],

  "references": [...],
  "citations":  [...]
}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

模板 [[目录]] 块写法（决定手动目录的格式，auto 模式也可以有）
--------------------------------------------------------------
[[目录]]
第一章 示例一级条目……1        ← 设好字体/加粗/缩进/前导符，这一行的格式 → TOC level 1
1.1 示例二级条目……1           ← 这一行的格式 → TOC level 2
1.1.1 示例三级条目……1         ← 这一行的格式 → TOC level 3
[[目录]]

三行各自的 pPr（缩进/tab/间距）和 rPr（字体/字号/加粗）会被分别提取。
如果只放两行，三级目录会复用二级样式；只放一行则全部复用该样式。
auto 模式下提取结果用于向 Word 样式表写入 TOC 1/2/3（如果不写，Word 会用自己的默认样式）。

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

依赖：
  pip install python-docx --break-system-packages
  pip install latex2mathml --break-system-packages   # formula/latex 可选
"""

from __future__ import annotations

from os import name
import re
import copy
import json
import io
import base64

import urllib.request
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt

# ── XML 命名空间 ───────────────────────────────────────────────
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M = "http://schemas.openxmlformats.org/officeDocument/2006/math"

TAG_RE  = re.compile(r"^\[\[(.+?)\]\]$")
CITE_RE = re.compile(r"\[(\d+)\]")

HEADING_TAGS = {"一级标题", "二级标题", "三级标题"}
_STATIC_TAGS = HEADING_TAGS | {"表格", "参考文献", "公式", "图片", "目录"}

TYPE_TO_TAG = {
    "heading1": "一级标题",
    "heading2": "二级标题",
    "heading3": "三级标题",
}
TAG_TO_STYLE_ID = {
    "一级标题": "Heading1",
    "二级标题": "Heading2",
    "三级标题": "Heading3",
}

_ALIGN_MAP = {
    "left":   WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right":  WD_ALIGN_PARAGRAPH.RIGHT,
}

SECTION_DEFAULT_TITLES = {
    "abstract":        "摘  要",
    "abstract_en":     "Abstract",
    "conclusion":      "结  论",
    "acknowledgement": "致  谢",
    "publications":    "已发表的学术论文目录",
    "custom":          "",
}

# ── 页脚预设样式 ───────────────────────────────────────────────
FOOTER_STYLES: dict[str, Optional[dict]] = {
    "roman_lower_center": {
        "fmt": "lowerRoman", "prefix": "",   "suffix": "",
        "instr": "PAGE",     "align": "center", "with_numpages": False,
    },
    "roman_upper_center": {
        "fmt": "upperRoman", "prefix": "",   "suffix": "",
        "instr": "PAGE",     "align": "center", "with_numpages": False,
    },
    "arabic_center": {
        "fmt": "decimal",    "prefix": "",   "suffix": "",
        "instr": "PAGE",     "align": "center", "with_numpages": False,
    },
    "arabic_dash": {
        "fmt": "decimal",    "prefix": "- ", "suffix": " -",
        "instr": "PAGE",     "align": "center", "with_numpages": False,
    },
    "arabic_page_x": {
        "fmt": "decimal",    "prefix": "第", "suffix": "页",
        "instr": "PAGE",     "align": "center", "with_numpages": False,
    },
    "arabic_slash": {
        "fmt": "decimal",    "prefix": "",   "suffix": "",
        "instr": "PAGE",     "align": "center", "with_numpages": True,
    },
    "none": None,
}

# LLM 系统提示（页脚解析）
_FOOTER_LLM_SYSTEM = """你是一个 Word 文档格式助手。
用户会用自然语言描述文档的页码需求。你需要把它映射到预设样式配置。

可用的 section 值：
  frontmatter  → 正文（第一个一级标题）出现之前的所有页
  mainmatter   → 正文开始（第一个一级标题）到文档末尾

可用的 style 值：
  roman_lower_center   小写罗马数字居中 (i, ii, iii...)
  roman_upper_center   大写罗马数字居中 (I, II, III...)
  arabic_center        阿拉伯数字居中 (1, 2, 3...)
  arabic_dash          -X- 格式居中 (-1-, -2-...)
  arabic_page_x        第X页格式居中 (第1页, 第2页...)
  arabic_slash         X/N 格式居中 (1/10, 2/10...)
  none                 无页脚

输出严格的 JSON 数组，不含任何解释：
[
  {"section": "frontmatter", "style": "roman_lower_center", "start": 1},
  {"section": "mainmatter",  "style": "arabic_dash",        "start": 1}
]

如果整个文档只有一种页码，section 填 mainmatter，frontmatter 用 none。
只输出 JSON，不要 markdown 代码块，不要任何其他文字。"""


# ══════════════════════════════════════════════════════════════
# 数据结构
# ══════════════════════════════════════════════════════════════

@dataclass
class TocLevelStyle:
    """单级目录条目的样式（pPr + rPr）"""
    pPr: object = None   # lxml Element
    rPr: object = None   # lxml Element


@dataclass
class StyleTemplate:
    """从模板 docx 中提取的各类格式样板"""
    table_proto:     object = None
    ref_pPr_proto:   object = None
    formula_pPr:     object = None
    image_pPr:       object = None
    caption_rPr:     object = None
    section_styles:  dict   = field(default_factory=dict)

    # 目录各级样式（最多三级），列表顺序 = [level1, level2, level3]
    # 如果模板里 [[目录]] 块只放了 N 行，则只有 N 个元素，其余由代码补齐
    toc_level_styles: list  = field(default_factory=list)   # List[TocLevelStyle]


# ══════════════════════════════════════════════════════════════
# 模板解析器
# ══════════════════════════════════════════════════════════════

def parse_template(doc_path: str,
                   extra_section_types: list[str] | None = None) -> StyleTemplate:
    doc  = Document(doc_path)
    body = doc.element.body
    st   = StyleTemplate()

    section_tags       = set(extra_section_types or [])
    section_title_tags = {f"{s}标题" for s in section_tags}
    all_tags = _STATIC_TAGS | section_tags | section_title_tags

    in_block: Optional[str] = None
    buf: list = []

    for child in body:
        local = child.tag.split("}")[-1]

        if local == "p":
            text = _para_text(child).strip()
            m    = TAG_RE.match(text)
            tag  = m.group(1) if m else None

            if tag in all_tags and in_block is None:
                in_block = tag
                buf = []
                continue

            if tag == in_block and in_block is not None:
                _extract_style(st, in_block, buf, section_tags)
                in_block = None
                buf = []
                continue

            if in_block is not None:
                buf.append(child)

        elif local in ("tbl", "oMathPara"):
            if in_block is not None:
                buf.append(child)

    return st


def _extract_style(st: StyleTemplate, tag: str, elems: list,
                   section_tags: set[str]):

    if tag == "表格":
        for e in elems:
            if e.tag.split("}")[-1] == "tbl":
                st.table_proto = copy.deepcopy(e)
                break

    elif tag == "参考文献":
        for e in elems:
            if e.tag.split("}")[-1] == "p":
                pPr = e.find(f"{{{W}}}pPr")
                if pPr is not None:
                    st.ref_pPr_proto = copy.deepcopy(pPr)
                    break

    elif tag == "公式":
        for e in elems:
            if e.tag.split("}")[-1] == "p":
                pPr = e.find(f"{{{W}}}pPr")
                if pPr is not None:
                    st.formula_pPr = copy.deepcopy(pPr)
                    break

    elif tag == "图片":
        for e in elems:
            if e.tag.split("}")[-1] == "p":
                pPr = e.find(f"{{{W}}}pPr")
                if pPr is not None and st.image_pPr is None:
                    st.image_pPr = copy.deepcopy(pPr)
                r = e.find(f"{{{W}}}r")
                if r is not None:
                    rPr = r.find(f"{{{W}}}rPr")
                    if rPr is not None and st.caption_rPr is None:
                        st.caption_rPr = copy.deepcopy(rPr)

    elif tag == "目录":
        # 提取 1/2/3 级目录条目样式
        # 块内每一个 <w:p> 对应一个级别，按顺序收集
        para_elems = [e for e in elems if e.tag.split("}")[-1] == "p"]
        for pe in para_elems[:3]:   # 最多取前三行
            pPr = pe.find(f"{{{W}}}pPr")
            rPr = pe.find(f".//{{{W}}}rPr")
            st.toc_level_styles.append(TocLevelStyle(
                pPr = copy.deepcopy(pPr) if pPr is not None else None,
                rPr = copy.deepcopy(rPr) if rPr is not None else None,
            ))

    elif tag in section_tags:
        ss = st.section_styles.setdefault(tag, {})
        for e in elems:
            if e.tag.split("}")[-1] == "p":
                pPr = e.find(f"{{{W}}}pPr")
                if pPr is not None:
                    ss["body_pPr"] = copy.deepcopy(pPr)
                    break

    elif tag.endswith("标题") and tag[:-2] in section_tags:
        stype = tag[:-2]
        ss = st.section_styles.setdefault(stype, {})
        for e in elems:
            if e.tag.split("}")[-1] == "p":
                pPr = e.find(f"{{{W}}}pPr")
                if pPr is not None:
                    ss["title_pPr"] = copy.deepcopy(pPr)
                r = e.find(f"{{{W}}}r")
                if r is not None:
                    rPr = r.find(f"{{{W}}}rPr")
                    if rPr is not None:
                        ss["title_rPr"] = copy.deepcopy(rPr)
                break


# ══════════════════════════════════════════════════════════════
# 目录渲染
# ══════════════════════════════════════════════════════════════

def _get_toc_level_style(st: StyleTemplate, level: int) -> TocLevelStyle:
    """
    获取第 level 级（1-based）的目录样式。
    若模板里不足三级，用已有的最后一级补齐；若完全没有，返回内置默认。
    """
    styles = st.toc_level_styles
    if styles:
        idx = min(level - 1, len(styles) - 1)
        return styles[idx]
    # 内置默认：逐级缩进 480 twips，带右对齐前导符 tab
    return _make_default_toc_style(level)


def _make_default_toc_style(level: int) -> TocLevelStyle:
    """当模板没有 [[目录]] 块时的内置兜底样式"""
    pPr = OxmlElement("w:pPr")

    # 缩进：一级0，二级480，三级960
    indent = (level - 1) * 480
    if indent > 0:
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(indent))
        pPr.append(ind)

    # 右对齐前导符 tab（约 A4 正文右边距）
    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"),    "right")
    tab.set(qn("w:pos"),    "9000")
    tab.set(qn("w:leader"), "dot")
    tabs.append(tab)
    pPr.append(tabs)

    # 间距：一级上方 6pt
    if level == 1:
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "120")
        pPr.append(spacing)

    rPr = OxmlElement("w:rPr")
    if level == 1:
        rPr.append(OxmlElement("w:b"))
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(24 - (level - 1) * 2))   # 12pt, 11pt, 10pt
    rPr.append(sz)

    return TocLevelStyle(pPr=pPr, rPr=rPr)


def _build_toc_entry_para(doc: Document, text: str, page_str: str,
                           style: TocLevelStyle) -> None:
    """
    向 doc 添加一个目录条目段落：
      [文字] [Tab→前导点→] [页码]
    """
    p = doc.add_paragraph()

    # 应用段落格式
    if style.pPr is not None:
        _apply_pPr(p._p, style.pPr)

    # 文字 run
    r_text = p.add_run(text)
    if style.rPr is not None:
        r_text._r.insert(0, copy.deepcopy(style.rPr))

    # Tab run（跳到前导符 tab stop）
    r_tab = p.add_run()
    tab_elem = OxmlElement("w:tab")
    r_tab._r.append(tab_elem)

    # 页码 run
    r_page = p.add_run(str(page_str))
    if style.rPr is not None:
        r_page._r.insert(0, copy.deepcopy(style.rPr))


def _render_toc_auto(doc: Document, item: dict, st: StyleTemplate) -> None:
    """
    AUTO 模式：插入 Word TOC 域。
    Word 打开时自动刷新，目录格式由文档 TOC 1/2/3 样式决定。
    如果模板提供了 [[目录]] 块，把样式写入文档的 TOC 1/2/3 样式定义。
    """
    toc_title      = item.get("title", "目  录")
    title_exclude  = item.get("toc_title_exclude", True)

    # 目录标题
    if toc_title:
        _add_heading_with_exclude(doc, toc_title, "Heading1",
                                  exclude=title_exclude)

    # 如果模板提供了样式，写入文档 TOC 1/2/3 样式
    if st.toc_level_styles:
        _inject_toc_styles_into_doc(doc, st)

    # 插入 TOC 域
    p = doc.add_paragraph()
    r = p.add_run()

    fc1 = OxmlElement("w:fldChar")
    fc1.set(qn("w:fldCharType"), "begin")
    fc1.set(qn("w:dirty"),       "true")   # 标记需要刷新
    r._r.append(fc1)

    r2 = p.add_run()
    it = OxmlElement("w:instrText")
    it.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    # \o "1-3"  扫描 Heading1-3
    # \h         条目带超链接（Ctrl+点击跳转）
    # \z         打印预览外隐藏页码和前导符
    # \u         使用段落 outline level（不仅限于内置 Heading 样式）
    it.text = r' TOC \o "1-3" \h \z \u '
    r2._r.append(it)

    r3 = p.add_run()
    fc2 = OxmlElement("w:fldChar")
    fc2.set(qn("w:fldCharType"), "end")
    r3._r.append(fc2)

    # 在 TOC 域后加一个空段落（Word 习惯）
    doc.add_paragraph()


def _inject_toc_styles_into_doc(doc: Document, st: StyleTemplate) -> None:
    """
    把从模板 [[目录]] 块提取到的 pPr/rPr 写入当前文档的
    TOC 1 / TOC 2 / TOC 3 段落样式定义。
    这样 Word 刷新 TOC 域时就会使用这些样式。
    """
    style_ids = ["TOC1", "TOC2", "TOC3"]
    styles_part = doc.part.styles

    for i, tls in enumerate(st.toc_level_styles[:3]):
        sid = style_ids[i]
        # 找到已有样式定义，或创建新的
        style_elem = styles_part.element.find(
            f".//{{{W}}}style[@{{{W}}}styleId='{sid}']"
        )
        if style_elem is None:
            # 创建最小样式定义
            style_elem = OxmlElement("w:style")
            style_elem.set(qn("w:type"),    "paragraph")
            style_elem.set(qn("w:styleId"), sid)
            name_e = OxmlElement("w:name")
            name_e.set(qn("w:val"), f"toc {i+1}")
            style_elem.append(name_e)
            styles_part.element.append(style_elem)

        # 写入 pPr
        if tls.pPr is not None:
            old_pPr = style_elem.find(f"{{{W}}}pPr")
            if old_pPr is not None:
                style_elem.remove(old_pPr)
            style_elem.append(copy.deepcopy(tls.pPr))

        # 写入 rPr（在 style 里叫 rPr，不是 pPr 的子元素）
        if tls.rPr is not None:
            old_rPr = style_elem.find(f"{{{W}}}rPr")
            if old_rPr is not None:
                style_elem.remove(old_rPr)
            style_elem.append(copy.deepcopy(tls.rPr))


def _render_toc_manual(doc: Document, item: dict,
                        st: StyleTemplate, toc_entries: list[dict]) -> None:
    """
    MANUAL 模式：用提供的 toc_entries 和模板样式手动排版目录。

    toc_entries 格式：
      [{"title": "摘  要",      "level": 1, "page": "i"},
       {"title": "第一章 引言", "level": 1, "page": 1},
       {"title": "1.1 研究背景","level": 2, "page": 2}, ...]

    优点：页码格式完全自定义（可以对前置页用罗马数字、正文用阿拉伯数字）。
    缺点：页码需要用户手动提供，或由外部流程（如 LibreOffice PDF 转换）计算。
    """
    toc_title     = item.get("title", "目  录")
    title_exclude = item.get("toc_title_exclude", True)

    if toc_title:
        _add_heading_with_exclude(doc, toc_title, "Heading1",
                                  exclude=title_exclude)

    for entry in toc_entries:
        level    = int(entry.get("level", 1))
        title    = entry.get("title", "")
        page_str = str(entry.get("page", ""))
        style    = _get_toc_level_style(st, level)
        _build_toc_entry_para(doc, title, page_str, style)

    doc.add_paragraph()


def _add_heading_with_exclude(doc: Document, text: str, style_id: str,
                               exclude: bool = False) -> None:
    """
    添加标题段落。
    exclude=True 时设置 outlineLevel=9，使该段落不被 TOC 域收录。
    """
    para = doc.add_paragraph()
    pPr  = para._p.get_or_add_pPr()

    ps = OxmlElement("w:pStyle")
    ps.set(qn("w:val"), style_id)
    pPr.insert(0, ps)

    if exclude:
        ol = OxmlElement("w:outlineLvl")
        ol.set(qn("w:val"), "9")   # 9 = 不进入任何大纲/目录
        pPr.append(ol)

    para.add_run(text)


# ══════════════════════════════════════════════════════════════
# 页脚系统（与 v4 完全一致，完整保留）
# ══════════════════════════════════════════════════════════════

def parse_footer_spec(spec: str, api_key: str | None = None) -> list[dict]:
    if api_key:
        return _parse_footer_via_llm(spec, api_key)
    return _parse_footer_heuristic(spec)


def _parse_footer_via_llm(spec: str, api_key: str) -> list[dict]:
    payload = json.dumps({
        "model": "claude-sonnet-4-20250514",
        "max_tokens": 512,
        "system": _FOOTER_LLM_SYSTEM,
        "messages": [{"role": "user", "content": spec}],
    }).encode()
    req = urllib.request.Request(
        "https://api.anthropic.com/v1/messages",
        data=payload,
        headers={
            "Content-Type":      "application/json",
            "x-api-key":         api_key,
            "anthropic-version": "2023-06-01",
        },
        method="POST",
    )
    try:
        with urllib.request.urlopen(req, timeout=15) as resp:
            body = json.loads(resp.read())
        text = body["content"][0]["text"].strip()
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text)
        config = json.loads(text)
        validated = []
        for it in config:
            if it.get("style") in FOOTER_STYLES and \
               it.get("section") in ("frontmatter", "mainmatter"):
                validated.append({
                    "section": it["section"],
                    "style":   it["style"],
                    "start":   int(it.get("start", 1)),
                })
        return validated
    except Exception as exc:
        print(f"⚠️  LLM 页脚解析失败（{exc}），退回规则匹配")
        return _parse_footer_heuristic(spec)


def _parse_footer_heuristic(spec: str) -> list[dict]:
    spec_l = spec.lower()

    def _detect_style(text: str) -> str:
        if "罗马" in text or "roman" in text.lower():
            return "roman_upper_center" if ("大写" in text or "upper" in text.lower()) \
                   else "roman_lower_center"
        if "-" in text and ("x" in text.lower() or "页" in text):
            return "arabic_dash"
        if "第" in text and "页" in text:
            return "arabic_page_x"
        if "/" in text:
            return "arabic_slash"
        if "无" in text or "none" in text.lower():
            return "none"
        return "arabic_center"

    front_kws = ["前置", "封面", "目录", "摘要", "前言", "frontmatter"]
    main_kws  = ["正文", "章节", "mainmatter", "阿拉伯"]
    result    = []

    if any(kw in spec_l for kw in front_kws):
        for kw in front_kws:
            idx = spec_l.find(kw)
            if idx != -1:
                snippet = spec[max(0, idx-5):idx+40]
                result.append({"section": "frontmatter",
                                "style": _detect_style(snippet), "start": 1})
                break

    if any(kw in spec_l for kw in main_kws):
        for kw in main_kws:
            idx = spec_l.find(kw)
            if idx != -1:
                snippet = spec[max(0, idx-5):idx+40]
                result.append({"section": "mainmatter",
                                "style": _detect_style(snippet), "start": 1})
                break

    if not result:
        result.append({"section": "mainmatter",
                        "style": _detect_style(spec), "start": 1})
    return result


def apply_footer_config(doc: Document, config: list[dict],
                        has_frontmatter: bool) -> None:
    front_cfg = next((c for c in config if c["section"] == "frontmatter"), None)
    main_cfg  = next((c for c in config if c["section"] == "mainmatter"),  None)
    sections  = doc.sections

    if has_frontmatter and len(sections) >= 2:
        if front_cfg:
            _write_section_footer(sections[0], front_cfg["style"],
                                  front_cfg.get("start", 1))
        if main_cfg:
            _write_section_footer(sections[1], main_cfg["style"],
                                  main_cfg.get("start", 1),
                                  unlink_from_prev=True)
    else:
        cfg = main_cfg or front_cfg
        if cfg:
            _write_section_footer(sections[-1], cfg["style"],
                                  cfg.get("start", 1))


def _write_section_footer(section, style_name: str, start_num: int = 1,
                           unlink_from_prev: bool = False) -> None:
    style  = FOOTER_STYLES.get(style_name)
    sectPr = section._sectPr
    for old in sectPr.findall(f"{{{W}}}pgNumType"):
        sectPr.remove(old)

    if style is None:
        footer = section.footer
        footer.is_linked_to_previous = False
        for para in footer.paragraphs:
            para.clear()
        return

    pgNumType = OxmlElement("w:pgNumType")
    pgNumType.set(qn("w:fmt"),   style["fmt"])
    pgNumType.set(qn("w:start"), str(start_num))
    sectPr.append(pgNumType)

    footer = section.footer
    if unlink_from_prev:
        footer.is_linked_to_previous = False
    for para in footer.paragraphs:
        para.clear()
    fp = footer.paragraphs[0]
    fp.alignment = _ALIGN_MAP.get(style.get("align", "center"),
                                  WD_ALIGN_PARAGRAPH.CENTER)

    if style.get("prefix"):
        fp.add_run(style["prefix"])
    _add_field_run(fp, style.get("instr", "PAGE"))
    if style.get("with_numpages"):
        fp.add_run("/")
        _add_field_run(fp, "NUMPAGES")
    if style.get("suffix"):
        fp.add_run(style["suffix"])


def _add_field_run(para, instr: str) -> None:
    r1 = para.add_run()
    fc1 = OxmlElement("w:fldChar")
    fc1.set(qn("w:fldCharType"), "begin")
    r1._r.append(fc1)

    r2 = para.add_run()
    it = OxmlElement("w:instrText")
    it.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    it.text = f" {instr} "
    r2._r.append(it)

    r3 = para.add_run()
    fc2 = OxmlElement("w:fldChar")
    fc2.set(qn("w:fldCharType"), "end")
    r3._r.append(fc2)


# ══════════════════════════════════════════════════════════════
# 通用章节块渲染
# ══════════════════════════════════════════════════════════════

def _render_section(doc: Document, item: dict, st: StyleTemplate) -> None:
    stype   = item.get("section_type", "custom")
    title   = item.get("title") or SECTION_DEFAULT_TITLES.get(stype, stype)
    value   = item.get("value", "")
    ss      = st.section_styles.get(stype, {})
    exclude = item.get("toc_exclude", False)

    # 标题
    title_para = doc.add_paragraph()
    title_pPr  = ss.get("title_pPr")
    if title_pPr is not None:
        _apply_pPr(title_para._p, title_pPr)
        if exclude:
            pPr = title_para._p.find(f"{{{W}}}pPr")
            if pPr is not None:
                ol = OxmlElement("w:outlineLvl")
                ol.set(qn("w:val"), "9")
                pPr.append(ol)
        r = title_para.add_run(title)
        title_rPr = ss.get("title_rPr")
        if title_rPr is not None:
            r._r.insert(0, copy.deepcopy(title_rPr))
    else:
        _add_heading_with_exclude(doc, title, "Heading1", exclude=exclude)
        # 已经通过 add_paragraph 创建了，把刚才 add_paragraph() 的空段删掉
        # （_add_heading_with_exclude 自己加段，title_para 是多余的）
        title_para._p.getparent().remove(title_para._p)

    # 正文
    body_pPr = ss.get("body_pPr")
    for para_text in _split_paragraphs(value):
        p = doc.add_paragraph()
        _apply_pPr(p._p, body_pPr)
        p.add_run(para_text)


def _split_paragraphs(text: str) -> list[str]:
    paras = re.split(r"\n{2,}", text)
    return [p.strip().replace("\n", " ") for p in paras if p.strip()]


# ══════════════════════════════════════════════════════════════
# 分节符插入
# ══════════════════════════════════════════════════════════════

def _insert_section_break(doc: Document) -> None:
    para = doc.add_paragraph()
    pPr  = para._p.get_or_add_pPr()
    sect = OxmlElement("w:sectPr")
    pgBrk = OxmlElement("w:type")
    pgBrk.set(qn("w:val"), "nextPage")
    sect.append(pgBrk)
    pPr.append(sect)


# ══════════════════════════════════════════════════════════════
# 生成器
# ══════════════════════════════════════════════════════════════

def generate(data: dict, st: StyleTemplate, output_path: str,
             footer_config: list[dict] | None = None) -> Document:
    """
    核心生成函数。

    文档装配顺序（固定）：
      前置章节（摘要等）
      [分节符，如果需要前/正文分页脚]
      目录（如果 content 里有 toc 块）
      正文内容
      后置章节（结论、致谢等）
      参考文献
    """
    toc_mode    = data.get("toc_mode", "auto")
    toc_entries = data.get("toc_entries", [])

    has_frontmatter = _has_frontmatter_sections(data.get("content", []))
    needs_split     = (has_frontmatter and footer_config is not None and
                       any(c["section"] == "frontmatter" for c in footer_config))

    doc = Document()

    split_inserted = False
    for item in data.get("content", []):
        t = item["type"]

        # 在第一个正文 heading1 前插入分节符
        if needs_split and not split_inserted and t == "heading1":
            _insert_section_break(doc)
            split_inserted = True

        if t in ("heading1", "heading2", "heading3"):
            exclude = item.get("toc_exclude", False)
            tag     = TYPE_TO_TAG[t]
            _add_heading_with_exclude(doc, item["value"],
                                      TAG_TO_STYLE_ID[tag], exclude=exclude)

        elif t == "body":
            doc.add_paragraph(item["value"])

        elif t == "table":
            _render_table(doc, item, st)

        elif t == "formula":
            _render_formula(doc, item, st)

        elif t == "image":
            _render_image(doc, item, st)

        elif t == "section":
            _render_section(doc, item, st)

        elif t == "toc":
            if toc_mode == "manual" and toc_entries:
                _render_toc_manual(doc, item, st, toc_entries)
            else:
                _render_toc_auto(doc, item, st)

    # 参考文献
    refs = data.get("references", [])
    if refs:
        _render_references(doc, refs, st)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    # 引用上标
    doc = Document(output_path)
    for cite in data.get("citations", []):
        insert_citation(doc, cite["ref_id"],
                        (cite["before"], cite.get("after", "")))

    # 页脚
    if footer_config:
        apply_footer_config(doc, footer_config,
                            has_frontmatter and split_inserted)

    doc.save(output_path)
    return doc


def _has_frontmatter_sections(content: list[dict]) -> bool:
    for item in content:
        if item["type"] == "heading1":
            return False
        if item["type"] == "section":
            return True
    return False


# ══════════════════════════════════════════════════════════════
# 原有渲染函数（完整保留）
# ══════════════════════════════════════════════════════════════

def _add_heading(doc: Document, text: str, style_id: str):
    """兼容旧接口（内部渲染参考文献标题等场景）"""
    _add_heading_with_exclude(doc, text, style_id, exclude=False)


def _render_table(doc: Document, item: dict, st: StyleTemplate):
    caption   = item.get("caption", "")
    rows_data = item["data"]
    if not rows_data:
        return
    num_cols = max(len(r) for r in rows_data)

    if caption:
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cap.runs:
            cap.runs[0].bold = True

    tbl = (_clone_table_with_data(st.table_proto, rows_data, num_cols)
           if st.table_proto is not None
           else _build_plain_table(rows_data, num_cols))

    body   = doc.element.body
    sectPr = body.find(f"{{{W}}}sectPr")
    if sectPr is not None:
        sectPr.addprevious(tbl)
    else:
        body.append(tbl)
    doc.add_paragraph("")


def _clone_table_with_data(proto, rows_data, num_cols):
    tbl = OxmlElement("w:tbl")
    proto_tblPr = proto.find(f"{{{W}}}tblPr")
    if proto_tblPr is not None:
        tbl.append(copy.deepcopy(proto_tblPr))
    proto_grid = proto.find(f"{{{W}}}tblGrid")
    if proto_grid is not None:
        grid = copy.deepcopy(proto_grid)
        cols = grid.findall(f"{{{W}}}gridCol")
        while len(cols) < num_cols:
            gc = OxmlElement("w:gridCol")
            gc.set(qn("w:w"), "1440")
            grid.append(gc)
            cols = grid.findall(f"{{{W}}}gridCol")
        tbl.append(grid)
    proto_rows   = proto.findall(f"{{{W}}}tr")
    header_proto = proto_rows[0] if proto_rows else None
    data_proto   = proto_rows[1] if len(proto_rows) > 1 else header_proto
    for row_idx, row_data in enumerate(rows_data):
        is_header = (row_idx == 0)
        tbl.append(_build_row_from_proto(
            header_proto if is_header else data_proto,
            row_data, num_cols, is_header))
    return tbl


def _build_row_from_proto(row_proto, row_data, num_cols, is_header):
    tr = OxmlElement("w:tr")
    if row_proto is not None:
        trPr = row_proto.find(f"{{{W}}}trPr")
        if trPr is not None:
            tr.append(copy.deepcopy(trPr))
    proto_cells = row_proto.findall(f"{{{W}}}tc") if row_proto is not None else []
    for col_idx in range(num_cols):
        cell_text  = row_data[col_idx] if col_idx < len(row_data) else ""
        cell_proto = (proto_cells[col_idx] if col_idx < len(proto_cells)
                      else (proto_cells[-1] if proto_cells else None))
        tr.append(_build_cell_from_proto(cell_proto, cell_text, is_header))
    return tr


def _build_cell_from_proto(cell_proto, text, is_header):
    tc = OxmlElement("w:tc")
    if cell_proto is not None:
        tcPr = cell_proto.find(f"{{{W}}}tcPr")
        if tcPr is not None:
            tc.append(copy.deepcopy(tcPr))
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    if cell_proto is not None:
        orig_p = cell_proto.find(f"{{{W}}}p")
        if orig_p is not None:
            pPr = orig_p.find(f"{{{W}}}pPr")
            if pPr is not None:
                p.append(copy.deepcopy(pPr))
            orig_r = orig_p.find(f"{{{W}}}r")
            if orig_r is not None:
                rPr = orig_r.find(f"{{{W}}}rPr")
                if rPr is not None:
                    new_rPr = copy.deepcopy(rPr)
                    if is_header and new_rPr.find(f"{{{W}}}b") is None:
                        new_rPr.insert(0, OxmlElement("w:b"))
                    r.append(new_rPr)
    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    r.append(t)
    p.append(r)
    tc.append(p)
    return tc


def _build_plain_table(rows_data, num_cols):
    total_w = 9360
    col_w   = total_w // num_cols
    col_ws  = [col_w] * num_cols
    col_ws[-1] += total_w - sum(col_ws)
    tbl = OxmlElement("w:tbl")
    tblPr = OxmlElement("w:tblPr")
    tblW  = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(total_w))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    tbl.append(tblPr)
    tblGrid = OxmlElement("w:tblGrid")
    for w in col_ws:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        tblGrid.append(gc)
    tbl.append(tblGrid)
    for row_idx, row_data in enumerate(rows_data):
        is_header = (row_idx == 0)
        tr = OxmlElement("w:tr")
        for col_idx in range(num_cols):
            cell_text = row_data[col_idx] if col_idx < len(row_data) else ""
            tc = OxmlElement("w:tc")
            tcPr = OxmlElement("w:tcPr")
            tcW  = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(col_ws[col_idx]))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)
            tc.append(tcPr)
            p = OxmlElement("w:p")
            r = OxmlElement("w:r")
            if is_header:
                rPr = OxmlElement("w:rPr")
                rPr.append(OxmlElement("w:b"))
                r.append(rPr)
            t = OxmlElement("w:t")
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t.text = cell_text
            r.append(t)
            p.append(r)
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)
    return tbl


def _render_references(doc: Document, refs: list[dict], st: StyleTemplate):
    _add_heading(doc, "参考文献", "Heading1")
    for ref in refs:
        para = doc.add_paragraph()
        pPr  = para._p.get_or_add_pPr()
        if st.ref_pPr_proto is not None:
            for child in st.ref_pPr_proto:
                if pPr.find(child.tag) is None:
                    pPr.append(copy.deepcopy(child))
        else:
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"),    "480")
            ind.set(qn("w:hanging"), "480")
            pPr.append(ind)
        run_num = para.add_run(f"[{ref['id']}] ")
        run_num.bold = True
        para.add_run(ref["text"])


# ══════════════════════════════════════════════════════════════
# 公式渲染（v3/v4 完整移植）
# ══════════════════════════════════════════════════════════════

def _render_formula(doc: Document, item: dict, st: StyleTemplate):
    label = item.get("label", "")
    if "omml" in item:
        omml_elem = _parse_omml_string(item["omml"])
    elif "latex" in item:
        omml_elem = _latex_to_omml(item["latex"])
    else:
        omml_elem = None

    body   = doc.element.body
    sectPr = body.find(f"{{{W}}}sectPr")

    if omml_elem is not None:
        p_elem = (_build_formula_para_with_label(omml_elem, label, st)
                  if label else _build_formula_para(omml_elem, st))
        if sectPr is not None:
            sectPr.addprevious(p_elem)
        else:
            body.append(p_elem)
    else:
        raw  = item.get("latex", item.get("omml", ""))
        txt  = f"{raw}    {label}" if label else raw
        para = doc.add_paragraph(txt)
        _apply_pPr(para._p, st.formula_pPr)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _build_formula_para(omml_elem, st):
    p = OxmlElement("w:p")
    p.append(_make_formula_pPr(st))
    p.append(copy.deepcopy(omml_elem))
    return p


def _build_formula_para_with_label(omml_elem, label, st):
    p = OxmlElement("w:p")
    pPr = _make_formula_pPr(st)
    tabs = OxmlElement("w:tabs")
    for val, pos in [("center", "4680"), ("right", "9360")]:
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), val)
        tab.set(qn("w:pos"), pos)
        tabs.append(tab)
    pPr.append(tabs)
    p.append(pPr)
    def tab_run():
        r = OxmlElement("w:r")
        r.append(OxmlElement("w:tab"))
        return r
    p.append(tab_run())
    p.append(copy.deepcopy(omml_elem))
    p.append(tab_run())
    r_label = OxmlElement("w:r")
    t_label = OxmlElement("w:t")
    t_label.text = label
    r_label.append(t_label)
    p.append(r_label)
    return p


def _make_formula_pPr(st):
    if st.formula_pPr is not None:
        return copy.deepcopy(st.formula_pPr)
    pPr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    pPr.append(jc)
    return pPr


def _parse_omml_string(omml_str):
    from lxml import etree
    try:
        return etree.fromstring(omml_str.encode())
    except Exception as exc:
        raise ValueError(f"无效的 OMML XML：{exc}") from exc


def _latex_to_omml(latex):
    try:
        import latex2mathml.converter
        from lxml import etree
    except ImportError:
        return None
    mathml_str  = latex2mathml.converter.convert(latex)
    mathml_elem = etree.fromstring(mathml_str.encode())
    return _mathml_to_omml_via_lxml(mathml_elem)


def _mathml_to_omml_via_lxml(mathml_elem):
    def _tag(e):
        return e.tag.split("}")[-1] if "}" in e.tag else e.tag

    def _text_run(text):
        r = OxmlElement("m:r")
        t = OxmlElement("m:t")
        t.text = text
        r.append(t)
        return r

    def _append_c(parent, elem):
        result = _conv(elem)
        if result is None: return
        if isinstance(result, list):
            for it in result:
                if it is not None: parent.append(it)
        else:
            parent.append(result)

    def _conv(elem):
        tag      = _tag(elem)
        children = list(elem)
        if tag in ("math", "mrow", "mstyle", "mpadded", "mphantom"):
            container = OxmlElement("m:oMath") if tag == "math" else None
            results   = [_conv(c) for c in children]
            if container is not None:
                for r in results:
                    if r is None: continue
                    (container.extend(r) if isinstance(r, list)
                     else container.append(r))
                return container
            return results
        elif tag == "mfrac":
            f = OxmlElement("m:f"); f.append(OxmlElement("m:fPr"))
            n = OxmlElement("m:num"); d = OxmlElement("m:den")
            if children: _append_c(n, children[0])
            if len(children)>1: _append_c(d, children[1])
            f.append(n); f.append(d); return f
        elif tag == "msup":
            ss = OxmlElement("m:sSup")
            e_ = OxmlElement("m:e"); s_ = OxmlElement("m:sup")
            if children: _append_c(e_, children[0])
            if len(children)>1: _append_c(s_, children[1])
            ss.append(e_); ss.append(s_); return ss
        elif tag == "msub":
            ss = OxmlElement("m:sSub")
            e_ = OxmlElement("m:e"); s_ = OxmlElement("m:sub")
            if children: _append_c(e_, children[0])
            if len(children)>1: _append_c(s_, children[1])
            ss.append(e_); ss.append(s_); return ss
        elif tag == "msqrt":
            rad = OxmlElement("m:rad")
            rp = OxmlElement("m:radPr"); dh = OxmlElement("m:degHide")
            dh.set(qn("m:val"), "1"); rp.append(dh); rad.append(rp)
            rad.append(OxmlElement("m:deg"))
            e_ = OxmlElement("m:e")
            for c in children: _append_c(e_, c)
            rad.append(e_); return rad
        elif tag in ("mn", "mi", "mo", "mtext", "ms"):
            text = (elem.text or "").strip()
            return _text_run(text) if text else None
        else:
            parts = [_conv(c) for c in children]
            parts = [p for p in parts if p is not None]
            if not parts and elem.text:
                return _text_run(elem.text.strip())
            if len(parts) == 1 and not isinstance(parts[0], list):
                return parts[0]
            wrap = OxmlElement("m:oMath")
            for part in parts:
                if isinstance(part, list):
                    for pp in part:
                        if pp is not None: wrap.append(pp)
                else:
                    wrap.append(part)
            return wrap

    oMath = _conv(mathml_elem)
    if oMath is None:
        oMath = OxmlElement("m:oMath")
        oMath.append(_text_run("?"))
    oMathPara = OxmlElement("m:oMathPara")
    oMathPara.append(oMath)
    return oMathPara


# ══════════════════════════════════════════════════════════════
# 图片渲染（v3/v4 完整移植）
# ══════════════════════════════════════════════════════════════

def _render_image(doc, item, st):
    width_in = float(item.get("width", 4.0))
    align    = _ALIGN_MAP.get(item.get("align", "center"), WD_ALIGN_PARAGRAPH.CENTER)
    caption  = item.get("caption", "")
    img_stream = _get_image_stream(item)
    if img_stream is None:
        ph = doc.add_paragraph(f"[图片占位：{caption or item.get('path','')}]")
        ph.alignment = align
        return
    para = doc.add_paragraph()
    para.alignment = align
    _apply_pPr(para._p, st.image_pPr)
    run = para.add_run()
    try:
        run.add_picture(img_stream, width=Inches(width_in))
    except Exception as exc:
        para.clear()
        para.add_run(f"[图片插入失败：{exc}]")
    if caption:
        cap_para = doc.add_paragraph()
        cap_para.alignment = align
        _apply_pPr(cap_para._p, st.image_pPr)
        r = cap_para.add_run(caption)
        if st.caption_rPr is not None:
            r._r.insert(0, copy.deepcopy(st.caption_rPr))


def _get_image_stream(item):
    if "path" in item:
        p = Path(item["path"])
        return io.BytesIO(p.read_bytes()) if p.exists() else None
    if "base64" in item:
        try:
            return io.BytesIO(base64.b64decode(item["base64"]))
        except Exception:
            return None
    return None


# ══════════════════════════════════════════════════════════════
# 引用插入
# ══════════════════════════════════════════════════════════════

def insert_citation(doc, ref_id, context):
    before, after = context
    for para in doc.paragraphs:
        full = para.text
        pos  = full.find(before)
        if pos == -1: continue
        insert_at = pos + len(before)
        if after and after[:5] not in full[insert_at:]: continue
        runs = para.runs
        if not runs: continue
        cur = 0
        target_idx, target_off = len(runs)-1, len(runs[-1].text)
        for ri, run in enumerate(runs):
            end = cur + len(run.text)
            if cur <= insert_at <= end:
                target_idx = ri; target_off = insert_at - cur; break
            cur = end
        target_run = runs[target_idx]
        orig_text  = target_run.text
        target_run.text = orig_text[:target_off]
        r_sup    = OxmlElement("w:r")
        orig_rPr = target_run._r.find(f"{{{W}}}rPr")
        new_rPr  = copy.deepcopy(orig_rPr) if orig_rPr is not None \
                   else OxmlElement("w:rPr")
        va = OxmlElement("w:vertAlign")
        va.set(qn("w:val"), "superscript")
        new_rPr.append(va)
        r_sup.append(new_rPr)
        t_sup = OxmlElement("w:t")
        t_sup.text = f"[{ref_id}]"
        r_sup.append(t_sup)
        r_tail = OxmlElement("w:r")
        if orig_rPr is not None:
            r_tail.append(copy.deepcopy(orig_rPr))
        t_tail = OxmlElement("w:t")
        t_tail.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t_tail.text = orig_text[target_off:]
        r_tail.append(t_tail)
        target_run._r.addnext(r_tail)
        target_run._r.addnext(r_sup)
        for run in para.runs:
            if run.text and f"[{ref_id}]" in run.text and not run.font.superscript:
                run.text = run.text.replace(f"[{ref_id}]", "")
        return True
    return False


# ══════════════════════════════════════════════════════════════
# 工具函数
# ══════════════════════════════════════════════════════════════

def _apply_pPr(p_elem, pPr_proto) -> None:
    if pPr_proto is None: return
    existing = p_elem.find(f"{{{W}}}pPr")
    if existing is None:
        existing = OxmlElement("w:pPr")
        p_elem.insert(0, existing)
    for child in pPr_proto:
        if existing.find(child.tag) is None:
            existing.append(copy.deepcopy(child))


def _para_text(p_elem) -> str:
    return "".join(t.text or "" for t in p_elem.iter(f"{{{W}}}t"))


# ══════════════════════════════════════════════════════════════
# 一站式接口
# ══════════════════════════════════════════════════════════════

def process(template_path: str, data_path: str, output_path: str,
            api_key: str | None = None, verbose: bool = True) -> Document:
    """
    一步到位：模板 + json → 输出 docx

    参数
    ----
    template_path : 含 [[标记]] 的格式模板 .docx
    data_path     : user_data.json 路径
    output_path   : 输出 .docx 路径
    api_key       : Anthropic API Key（LLM 解析页脚描述，可不填）
    """
    data = json.loads(Path(data_path).read_text(encoding="utf-8"))

    section_types = list({
        item["section_type"]
        for item in data.get("content", [])
        if item.get("type") == "section"
    })

    st = parse_template(template_path, extra_section_types=section_types)

    # 页脚配置
    footer_config = None
    if "page_footer_config" in data:
        footer_config = data["page_footer_config"]
        if verbose:
            print(f"📑  页脚配置（直接）：{footer_config}")
    elif "page_footer_spec" in data:
        if verbose:
            print(f"🤖  解析页脚描述：{data['page_footer_spec']!r}")
        footer_config = parse_footer_spec(data["page_footer_spec"], api_key)
        if verbose:
            print(f"   → 解析结果：{footer_config}")

    if verbose:
        content      = data.get("content", [])
        n_section    = sum(1 for x in content if x.get("type") == "section")
        n_toc        = sum(1 for x in content if x.get("type") == "toc")
        toc_mode     = data.get("toc_mode", "auto")
        toc_styles_n = len(st.toc_level_styles)
        print(f"📐  模板：表格={'有' if st.table_proto else '无'}  "
              f"文献pPr={'有' if st.ref_pPr_proto else '无'}  "
              f"章节样式={list(st.section_styles.keys()) or '无'}  "
              f"TOC级样式={toc_styles_n}级")
        print(f"📋  内容：{len(content)} 块（{n_section} 章节块，{n_toc} 目录块，"
              f"TOC模式={toc_mode}）  "
              f"文献={len(data.get('references',[]))}  "
              f"引用={len(data.get('citations',[]))}")

    doc = generate(data, st, output_path, footer_config=footer_config)

    if verbose:
        print(f"✅  已生成：{output_path}")
    return doc

if __name__ == "__main__":
    process(template_path='data/full_template.docx',data_path='data/full_user_data.json',output_path='data/output.docx')
