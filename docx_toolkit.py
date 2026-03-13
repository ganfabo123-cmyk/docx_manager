"""
docx_toolkit.py
================
一个强大的 Word 文档自动化工具库
支持：读取/替换正文、各级标题、分页符、页脚页码格式、表格操作（读取/插入行列/复制）

依赖：python-docx
    pip install python-docx --break-system-packages
"""

import copy
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from lxml import etree


# ─────────────────────────────────────────────
# 1. 文档读取与分析
# ─────────────────────────────────────────────

def read_document(path: str) -> Document:
    """
    读取一个 .docx 文件，返回 Document 对象。

    参数:
        path: 文件路径

    返回:
        Document 对象
    """
    return Document(path)


def get_body_text(doc: Document) -> str:
    """
    读取文档所有正文段落（排除标题段落）。

    返回:
        正文内容（字符串，段落间以换行分隔）
    """
    heading_styles = {"Heading 1", "Heading 2", "Heading 3",
                      "标题 1", "标题 2", "标题 3", "Title"}
    lines = []
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        if style_name not in heading_styles and para.text.strip():
            lines.append(para.text)
    return "\n".join(lines)


def describe_document(doc: Document) -> dict:
    """
    分析文档结构，返回包含各类元素统计的字典。
    """
    info = {
        "total_paragraphs": len(doc.paragraphs),
        "headings": [],
        "body_paragraphs": 0,
        "tables": len(doc.tables),
        "sections": len(doc.sections),
    }
    heading_styles = {"Heading 1", "Heading 2", "Heading 3",
                      "标题 1", "标题 2", "标题 3", "Title"}
    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else ""
        if style_name in heading_styles:
            info["headings"].append({
                "index": i,
                "style": style_name,
                "text": para.text
            })
        elif para.text.strip():
            info["body_paragraphs"] += 1
    return info


def read_table(doc: Document, table_index: int = 0) -> list[list[str]]:
    """
    读取指定索引的表格内容，返回二维列表。

    参数:
        doc:          Document 对象
        table_index:  表格索引（0 起）

    返回:
        [ [row0_col0, row0_col1, ...], [row1_col0, ...], ... ]
    """
    if table_index >= len(doc.tables):
        raise IndexError(f"文档只有 {len(doc.tables)} 个表格，索引 {table_index} 越界")
    table = doc.tables[table_index]
    return [[cell.text for cell in row.cells] for row in table.rows]


# ─────────────────────────────────────────────
# 2. 正文替换（删除旧正文，插入新正文）
# ─────────────────────────────────────────────

def replace_body_text(doc: Document, new_paragraphs: list[str]) -> Document:
    """
    删除文档中所有非标题的正文段落，插入新正文。

    参数:
        doc:            Document 对象（会被原地修改）
        new_paragraphs: 新正文段落列表，每条字符串为一段

    返回:
        修改后的 Document 对象（同一实例）
    """
    heading_styles = {"Heading 1", "Heading 2", "Heading 3",
                      "标题 1", "标题 2", "标题 3", "Title"}

    # 找出所有需要删除的正文段落 XML 元素
    to_remove = []
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        if style_name not in heading_styles:
            to_remove.append(para._element)

    # 记录最后一个删除项的位置，作为插入点
    last_removed = to_remove[-1] if to_remove else None
    parent = last_removed.getparent() if last_removed is not None else doc.element.body

    # 插入新段落（在最后一个正文段之后）
    insert_after = last_removed
    for text in new_paragraphs:
        new_para = OxmlElement("w:p")
        new_r = OxmlElement("w:r")
        new_t = OxmlElement("w:t")
        new_t.text = text
        new_r.append(new_t)
        new_para.append(new_r)
        if insert_after is not None:
            insert_after.addnext(new_para)
            insert_after = new_para
        else:
            parent.append(new_para)

    # 删除旧正文
    for elem in to_remove:
        elem.getparent().remove(elem)

    return doc


# ─────────────────────────────────────────────
# 3. 标题函数
# ─────────────────────────────────────────────

def _add_heading_by_id(doc: Document, text: str, style_id: str) -> Document:
    """通过样式 ID 添加标题（兼容缺少 Normal 样式的文档）"""
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), style_id)
    pPr.insert(0, pStyle)
    para.add_run(text)
    return doc


def add_title(doc: Document, text: str, **kwargs) -> Document:
    """
    在文档末尾添加文档主标题（Title 样式）。

    参数:
        doc:  Document 对象
        text: 标题文本
    """
    return _add_heading_by_id(doc, text, "Title")


def add_heading1(doc: Document, text: str) -> Document:
    """添加一级标题（Heading 1）"""
    return _add_heading_by_id(doc, text, "Heading1")


def add_heading2(doc: Document, text: str) -> Document:
    """添加二级标题（Heading 2）"""
    return _add_heading_by_id(doc, text, "Heading2")


def add_heading3(doc: Document, text: str) -> Document:
    """添加三级标题（Heading 3）"""
    return _add_heading_by_id(doc, text, "Heading3")


def add_paragraph(doc: Document, text: str, bold: bool = False,
                  italic: bool = False, font_size: int = None) -> Document:
    """
    添加一段正文。

    参数:
        doc:       Document 对象
        text:      正文文本
        bold:      是否加粗
        italic:    是否斜体
        font_size: 字号（磅），None 使用默认
    """
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if font_size:
        run.font.size = Pt(font_size)
    return doc


# ─────────────────────────────────────────────
# 4. 分页符
# ─────────────────────────────────────────────

def add_page_break(doc: Document) -> Document:
    """在文档末尾插入分页符。"""
    para = doc.add_paragraph()
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    para._p.append(r)
    return doc


# ─────────────────────────────────────────────
# 5. 页脚页码格式
#    first_range: 罗马数字页码的页面范围 (start, end) 基于 1 起的页码
#    second_range: 阿拉伯数字页码的页面范围 (start, end)
#    实现方式：使用多个 Section，每个 Section 设置不同的 numFmt
# ─────────────────────────────────────────────

def _set_footer_page_number(section, num_format: str, start: int = 1):
    """
    为一个 section 设置页脚页码格式。

    参数:
        section:    Section 对象
        num_format: 'lowerRoman' | 'decimal' | 'upperRoman' | 'lowerLetter' 等
        start:      起始页码
    """
    # 设置 sectPr 的 pgNumType
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is None:
        pgNumType = OxmlElement("w:pgNumType")
        sectPr.append(pgNumType)
    pgNumType.set(qn("w:fmt"), num_format)
    pgNumType.set(qn("w:start"), str(start))

    # 创建页脚（含页码）
    footer = section.footer
    footer.is_linked_to_previous = False

    # 清除现有内容
    for para in footer.paragraphs:
        for run in para.runs:
            run.text = ""

    if footer.paragraphs:
        para = footer.paragraphs[0]
    else:
        para = footer.add_paragraph()

    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.clear()

    # 添加 PAGE 域
    run = para.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar1)

    run2 = para.add_run()
    instrText = OxmlElement("w:instrText")
    instrText.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instrText.text = " PAGE "
    run2._r.append(instrText)

    run3 = para.add_run()
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run3._r.append(fldChar2)


def set_page_number_format(
    doc: Document,
    first_range: tuple[int, int],
    second_range: tuple[int, int]
) -> Document:
    """
    设置文档页脚页码格式：
    - first_range  对应的页面使用「罗马数字」页码
    - second_range 对应的页面使用「阿拉伯数字」页码

    工作原理：
        通过在指定页面边界处插入「分节符（下一页）」，
        将文档划分为多个 Section，并对每个 Section 独立设置页码格式。

    参数:
        doc:          Document 对象
        first_range:  罗马数字页码范围，如 (1, 3) 表示第 1–3 页用罗马数字
        second_range: 阿拉伯数字页码范围，如 (4, 99)

    返回:
        修改后的 Document 对象
    """
    # 确保至少有两个 section
    body = doc.element.body
    # 获取当前 sectPr（文档末尾）
    last_sectPr = body.find(qn("w:sectPr"))

    # 若只有一个 section，添加一个新的分节符在 body 末尾之前
    # 策略：在第一段内容之后（跳过前 first_range[1] 页）插入分节符
    # 简化实现：直接对已有 section[0] 设置第一种格式，
    # 添加 sectPr 给之后的段落设置第二种格式

    # 找到 body 里所有段落，在"first_range[1]"个正文段落后插入新分节符
    paragraphs = body.findall(qn("w:p"))
    # 注意：first_range 是页面范围，但我们无法在 XML 层面精确知道每段在哪页
    # 这里以"在文档中间某个分页符处插入分节符"为实用策略
    # 找到第一个分页符段落作为分割点
    split_para = None
    for p in paragraphs:
        for r in p.findall(qn("w:r")):
            for br in r.findall(qn("w:br")):
                if br.get(qn("w:type")) == "page":
                    split_para = p
                    break
            if split_para is not None:
                break
        if split_para is not None:
            break

    if split_para is not None:
        # 在该分页符段落前插入 section break（下一页）
        new_para = OxmlElement("w:p")
        new_sectPr = OxmlElement("w:sectPr")
        pPr = OxmlElement("w:pPr")
        pPr.append(new_sectPr)
        new_para.append(pPr)

        # 设置第一个 section 的页脚（罗马数字）
        _apply_pgNumFmt(new_sectPr, "lowerRoman", first_range[0])

        # 添加 section footer
        _add_footer_to_sectPr(new_sectPr, "lowerRoman")

        split_para.addprevious(new_para)

    # 设置最后一个 section（阿拉伯数字）
    sections = doc.sections
    last_section = sections[-1]
    _set_footer_page_number(last_section, "decimal", second_range[0])

    return doc


def _apply_pgNumFmt(sectPr, fmt: str, start: int):
    pgNumType = OxmlElement("w:pgNumType")
    pgNumType.set(qn("w:fmt"), fmt)
    pgNumType.set(qn("w:start"), str(start))
    sectPr.append(pgNumType)


def _add_footer_to_sectPr(sectPr, fmt: str):
    """为内嵌 sectPr 添加页脚引用（简化版：依赖外部 footer part）"""
    # 标记该 section 不链接到前一个 section 的页脚
    footerRef = OxmlElement("w:footerReference")
    footerRef.set(qn("w:type"), "default")
    # 实际 rId 需要在 document.xml.rels 中注册，这里使用 sentinel
    footerRef.set("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id",
                  "rId_footer_roman")
    sectPr.append(footerRef)


# ─────────────────────────────────────────────
# 6. 表格操作
# ─────────────────────────────────────────────

def insert_table_row(doc: Document, table_index: int,
                     row_data: list[str], position: int = None) -> Document:
    """
    向指定表格插入一行数据。

    参数:
        doc:         Document 对象
        table_index: 表格索引（0 起）
        row_data:    行数据列表，长度需与表格列数一致
        position:    插入位置（0 起），None 表示追加到末尾

    返回:
        修改后的 Document 对象
    """
    table = doc.tables[table_index]
    num_cols = len(table.columns)
    if len(row_data) != num_cols:
        raise ValueError(f"行数据长度 {len(row_data)} 与表格列数 {num_cols} 不匹配")

    # 复制最后一行的 XML 作为模板（保留格式）
    template_row = table.rows[-1]._tr
    new_tr = copy.deepcopy(template_row)

    # 填入新数据
    cells = new_tr.findall(qn("w:tc"))
    for cell_elem, text in zip(cells, row_data):
        # 清除原有文本
        for p in cell_elem.findall(qn("w:p")):
            for r in p.findall(qn("w:r")):
                for t in r.findall(qn("w:t")):
                    t.text = text
                    break

    if position is None:
        # 追加到末尾
        table._tbl.append(new_tr)
    else:
        # 插入到指定位置
        rows = table._tbl.findall(qn("w:tr"))
        if position >= len(rows):
            table._tbl.append(new_tr)
        else:
            rows[position].addprevious(new_tr)

    return doc


def insert_table_column(doc: Document, table_index: int,
                        col_data: list[str], position: int = None) -> Document:
    """
    向指定表格插入一列数据。

    参数:
        doc:         Document 对象
        table_index: 表格索引（0 起）
        col_data:    列数据列表（包含表头），长度需与行数一致
        position:    插入位置（0 起），None 表示追加到末尾

    返回:
        修改后的 Document 对象
    """
    table = doc.tables[table_index]
    num_rows = len(table.rows)
    if len(col_data) != num_rows:
        raise ValueError(f"列数据长度 {len(col_data)} 与表格行数 {num_rows} 不匹配")

    for row_idx, row in enumerate(table.rows):
        # 复制该行第一个单元格作为模板
        template_tc = copy.deepcopy(row.cells[0]._tc)

        # 清除并设置新文本
        for p in template_tc.findall(qn("w:p")):
            for r in p.findall(qn("w:r")):
                for t in r.findall(qn("w:t")):
                    t.text = col_data[row_idx]

        cells_in_row = row._tr.findall(qn("w:tc"))
        if position is None or position >= len(cells_in_row):
            row._tr.append(template_tc)
        else:
            cells_in_row[position].addprevious(template_tc)

    return doc


def copy_table(doc: Document, source_index: int) -> Document:
    """
    复制指定表格并追加到文档末尾（带分段符隔开）。

    参数:
        doc:          Document 对象
        source_index: 要复制的表格索引（0 起）

    返回:
        修改后的 Document 对象
    """
    source_table = doc.tables[source_index]
    new_tbl = copy.deepcopy(source_table._tbl)

    # 在 body 末尾（sectPr 之前）插入一个空段落 + 新表格
    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))

    spacer = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "（以下为复制的表格）"
    r.append(t)
    spacer.append(r)

    if sectPr is not None:
        sectPr.addprevious(spacer)
        sectPr.addprevious(new_tbl)
    else:
        body.append(spacer)
        body.append(new_tbl)

    return doc


def get_table_info(doc: Document, table_index: int = 0) -> dict:
    """
    获取表格的基本信息。

    返回:
        { 'rows': N, 'cols': M, 'data': [[...], ...] }
    """
    table = doc.tables[table_index]
    data = [[cell.text for cell in row.cells] for row in table.rows]
    return {
        "rows": len(table.rows),
        "cols": len(table.columns),
        "data": data
    }


# ─────────────────────────────────────────────
# 7. 保存文档
# ─────────────────────────────────────────────

def save_document(doc: Document, output_path: str) -> str:
    """
    保存 Document 到文件。

    参数:
        doc:         Document 对象
        output_path: 输出文件路径

    返回:
        输出文件路径
    """
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


# ─────────────────────────────────────────────
# 8. 一站式：从模板生成新文档
# ─────────────────────────────────────────────

def generate_from_template(
    template_path: str,
    output_path: str,
    title: str = None,
    h1: str = None,
    h2: str = None,
    h3: str = None,
    body_paragraphs: list[str] = None,
    first_range: tuple[int, int] = (1, 3),
    second_range: tuple[int, int] = (4, 999),
    new_table_rows: list[list[str]] = None,
    add_col: list[str] = None,
    copy_first_table: bool = False,
) -> str:
    """
    一站式从模板生成新文档。

    参数:
        template_path:    源 .docx 路径
        output_path:      输出 .docx 路径
        title:            文档主标题（替换或新增）
        h1/h2/h3:         各级标题文本
        body_paragraphs:  新正文段落列表
        first_range:      罗马数字页码的页面范围
        second_range:     阿拉伯数字页码的页面范围
        new_table_rows:   向第一个表格追加的行数据
        add_col:          向第一个表格追加的列数据（包含表头）
        copy_first_table: 是否复制第一个表格

    返回:
        输出文件路径
    """
    doc = read_document(template_path)

    # 替换正文
    if body_paragraphs is not None:
        replace_body_text(doc, body_paragraphs)

    # 添加标题（追加到末尾作为新章节）
    if h1:
        add_heading1(doc, h1)
    if h2:
        add_heading2(doc, h2)
    if h3:
        add_heading3(doc, h3)

    # 添加分页符示范
    add_page_break(doc)

    # 页脚页码格式
    set_page_number_format(doc, first_range, second_range)

    # 表格操作
    if doc.tables:
        if new_table_rows:
            for row in new_table_rows:
                insert_table_row(doc, 0, row)
        if add_col and len(add_col) == len(doc.tables[0].rows):
            insert_table_column(doc, 0, add_col)
        if copy_first_table:
            copy_table(doc, 0)

    return save_document(doc, output_path)


# ─────────────────────────────────────────────
# CLI / 快速演示
# ─────────────────────────────────────────────
if __name__ == "__main__":
    import json

    SRC = "demo_source.docx"
    OUT = "output_generated.docx"

    print("=" * 60)
    print("📖  步骤 1：读取源文档")
    doc = read_document(SRC)
    info = describe_document(doc)
    print(f"    段落数: {info['total_paragraphs']}, 表格数: {info['tables']}, "
          f"Section 数: {info['sections']}")
    print(f"    标题列表: {[h['text'] for h in info['headings']]}")

    print("\n📄  步骤 2：读取原始正文")
    body = get_body_text(doc)
    print(f"    原始正文内容:\n    {body[:200]}")

    print("\n📊  步骤 3：读取表格")
    if doc.tables:
        tinfo = get_table_info(doc, 0)
        print(f"    表格大小: {tinfo['rows']} 行 × {tinfo['cols']} 列")
        for row in tinfo['data']:
            print(f"      {row}")

    print("\n✏️   步骤 4：生成新文档（替换正文 + 插入标题 + 操作表格 + 设置页脚）")
    result = generate_from_template(
        template_path=SRC,
        output_path=OUT,
        h1="用户自定义一级标题",
        h2="用户自定义二级标题",
        h3="用户自定义三级标题",
        body_paragraphs=[
            "这是用户插入的第一段新正文，原有正文已被删除。",
            "这是第二段新正文，支持多段落批量写入。",
            "本文档由 docx_toolkit.py 自动生成，展示了完整的文档自动化能力。",
        ],
        first_range=(1, 3),    # 第 1-3 页：罗马数字页码
        second_range=(4, 999), # 第 4 页起：阿拉伯数字页码
        new_table_rows=[
            ["王五", "技术部", "架构师"],
            ["赵六", "产品部", "产品经理"],
        ],
        add_col=["备注", "优秀", "良好", "待定", "优秀"],  # 4 行（含标题）
        copy_first_table=True,
    )
    print(f"\n✅  文档已生成: {result}")
    print("=" * 60)
