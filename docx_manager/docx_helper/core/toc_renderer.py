from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import copy

from .constants import W, TAG_TO_STYLE_ID, TYPE_TO_TAG
from ..models.models import StyleTemplate, TocLevelStyle


def _get_toc_level_style(st: StyleTemplate, level: int) -> TocLevelStyle:
    """
    获取第 level 级（1-based）的目录样式。
    若模板里不足三级，用已有的最后一级补齐；若完全没有，返回内置默认。
    """
    styles = st.toc_level_styles
    if styles:
        idx = min(level - 1, len(styles) - 1)
        return styles[idx]
    # 内置默认：逐级缩进 480 twips，带右对齐前导符 tab
    return _make_default_toc_style(level)


def _make_default_toc_style(level: int) -> TocLevelStyle:
    """当模板没有 [[目录]] 块时的内置兜底样式"""
    pPr = OxmlElement("w:pPr")

    # 缩进：一级0，二级480，三级960
    indent = (level - 1) * 480
    if indent > 0:
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(indent))
        pPr.append(ind)

    # 右对齐前导符 tab（约 A4 正文右边距）
    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"),    "right")
    tab.set(qn("w:pos"),    "9000")
    tab.set(qn("w:leader"), "dot")
    tabs.append(tab)
    pPr.append(tabs)

    # 间距：一级上方 6pt
    if level == 1:
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "120")
        pPr.append(spacing)

    rPr = OxmlElement("w:rPr")
    if level == 1:
        rPr.append(OxmlElement("w:b"))
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(24 - (level - 1) * 2))   # 12pt, 11pt, 10pt
    rPr.append(sz)

    return TocLevelStyle(pPr=pPr, rPr=rPr)


def _build_toc_entry_para(doc: Document, text: str, page_str: str, style: TocLevelStyle) -> None:
    """
    向 doc 添加一个目录条目段落：
      [文字] [Tab→前导点→] [页码]
    """
    p = doc.add_paragraph()

    # 应用段落格式
    if style.pPr is not None:
        _apply_pPr(p._p, style.pPr)

    # 文字 run
    r_text = p.add_run(text)
    if style.rPr is not None:
        r_text._r.insert(0, copy.deepcopy(style.rPr))

    # Tab run（跳到前导符 tab stop）
    r_tab = p.add_run()
    tab_elem = OxmlElement("w:tab")
    r_tab._r.append(tab_elem)

    # 页码 run
    r_page = p.add_run(str(page_str))
    if style.rPr is not None:
        r_page._r.insert(0, copy.deepcopy(style.rPr))


def _render_toc_auto(doc: Document, item: dict, st: StyleTemplate) -> None:
    """
    AUTO 模式：插入 Word TOC 域。
    Word 打开时自动刷新，目录格式由文档 TOC 1/2/3 样式决定。
    如果模板提供了 [[目录]] 块，把样式写入文档的 TOC 1/2/3 样式定义。
    """
    toc_title      = item.get("title", "目  录")
    title_exclude  = item.get("toc_title_exclude", True)

    # 目录标题（格式来自模板 [[一级标题]] 块）
    if toc_title:
        _add_heading_with_exclude(doc, toc_title, "heading1", exclude=title_exclude, st=st)

    # 如果模板提供了样式，写入文档 TOC 1/2/3 样式
    if st.toc_level_styles:
        _inject_toc_styles_into_doc(doc, st)

    # 插入 TOC 域
    p = doc.add_paragraph()
    r = p.add_run()

    fc1 = OxmlElement("w:fldChar")
    fc1.set(qn("w:fldCharType"), "begin")
    fc1.set(qn("w:dirty"),       "true")   # 标记需要刷新
    r._r.append(fc1)

    r2 = p.add_run()
    it = OxmlElement("w:instrText")
    it.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    # \o "1-3"  扫描 Heading1-3
    # \h         条目带超链接（Ctrl+点击跳转）
    # \z         打印预览外隐藏页码和前导符
    # \u         使用段落 outline level（不仅限于内置 Heading 样式）
    it.text = r' TOC \o "1-3" \h \z \u '
    r2._r.append(it)

    r3 = p.add_run()
    fc2 = OxmlElement("w:fldChar")
    fc2.set(qn("w:fldCharType"), "end")
    r3._r.append(fc2)

    # 在 TOC 域后加一个空段落（Word 习惯）
    doc.add_paragraph()


def _inject_toc_styles_into_doc(doc: Document, st: StyleTemplate) -> None:
    """
    把从模板 [[目录]] 块提取到的 pPr/rPr 写入当前文档的
    TOC 1 / TOC 2 / TOC 3 段落样式定义。
    这样 Word 刷新 TOC 域时就会使用这些样式。
    """
    style_ids = ["TOC1", "TOC2", "TOC3"]
    styles_part = doc.part.styles

    for i, tls in enumerate(st.toc_level_styles[:3]):
        sid = style_ids[i]
        # 找到已有样式定义，或创建新的
        style_elem = styles_part.element.find(
            f".//{{{W}}}style[@{{{W}}}styleId='{sid}']"
        )
        if style_elem is None:
            # 创建最小样式定义
            style_elem = OxmlElement("w:style")
            style_elem.set(qn("w:type"),    "paragraph")
            style_elem.set(qn("w:styleId"), sid)
            name_e = OxmlElement("w:name")
            name_e.set(qn("w:val"), f"toc {i+1}")
            style_elem.append(name_e)
            styles_part.element.append(style_elem)

        # 写入 pPr
        if tls.pPr is not None:
            old_pPr = style_elem.find(f"{{{W}}}pPr")
            if old_pPr is not None:
                style_elem.remove(old_pPr)
            style_elem.append(copy.deepcopy(tls.pPr))

        # 写入 rPr（在 style 里叫 rPr，不是 pPr 的子元素）
        if tls.rPr is not None:
            old_rPr = style_elem.find(f"{{{W}}}rPr")
            if old_rPr is not None:
                style_elem.remove(old_rPr)
            style_elem.append(copy.deepcopy(tls.rPr))


def _render_toc_manual(doc: Document, item: dict, st: StyleTemplate, toc_entries: list[dict]) -> None:
    """
    MANUAL 模式：用提供的 toc_entries 和模板样式手动排版目录。

    toc_entries 格式：
      [{"title": "摘  要",      "level": 1, "page": "i"},
       {"title": "第一章 引言", "level": 1, "page": 1},
       {"title": "1.1 研究背景","level": 2, "page": 2}, ...]

    优点：页码格式完全自定义（可以对前置页用罗马数字、正文用阿拉伯数字）。
    缺点：页码需要用户手动提供，或由外部流程（如 LibreOffice PDF 转换）计算。
    """
    toc_title     = item.get("title", "目  录")
    title_exclude = item.get("toc_title_exclude", True)

    if toc_title:
        _add_heading_with_exclude(doc, toc_title, "heading1", exclude=title_exclude, st=st)

    for entry in toc_entries:
        level    = int(entry.get("level", 1))
        title    = entry.get("title", "")
        page_str = str(entry.get("page", ""))
        style    = _get_toc_level_style(st, level)
        _build_toc_entry_para(doc, title, page_str, style)

    doc.add_paragraph()


def _add_heading_with_exclude(doc: Document, text: str, style_key: str, exclude: bool = False, st: StyleTemplate | None = None) -> None:
    """
    添加标题段落。格式优先级：
      1. 模板 [[一/二/三级标题]] 块提取的 pPr + rPr
      2. Word 内置 Heading1/2/3 样式（兜底）

    exclude=True 时写入 outlineLevel=9，不被 TOC 域收录。
    style_key: "heading1" / "heading2" / "heading3"（或直接传 Word styleId 作兜底）
    """
    para = doc.add_paragraph()
    p    = para._p
    pPr  = p.get_or_add_pPr()

    # 尝试从模板提取样式
    hs = (st.heading_styles.get(style_key) if st else None)

    if hs is not None and (hs.pPr is not None or hs.rPr is not None):
        # 完整替换 pPr（格式由模板决定）
        if hs.pPr is not None:
            p.remove(pPr)
            new_pPr = copy.deepcopy(hs.pPr)
            # 保留或添加 pStyle，确保标题样式正确
            builtin_id = TAG_TO_STYLE_ID.get(TYPE_TO_TAG.get(style_key, ""), "Heading1")
            pStyle = new_pPr.find(f"{{{W}}}pStyle")
            if pStyle is None:
                pStyle = OxmlElement("w:pStyle")
                pStyle.set(qn("w:val"), builtin_id)
                new_pPr.insert(0, pStyle)
            else:
                pStyle.set(qn("w:val"), builtin_id)
            p.insert(0, new_pPr)
            pPr = new_pPr
    else:
        # 兜底：使用 Word 内置 Heading 样式
        builtin_id = TAG_TO_STYLE_ID.get(TYPE_TO_TAG.get(style_key, ""), "Heading1")
        ps = OxmlElement("w:pStyle")
        ps.set(qn("w:val"), builtin_id)
        pPr.insert(0, ps)

    if exclude:
        ol = OxmlElement("w:outlineLvl")
        ol.set(qn("w:val"), "9")
        pPr.append(ol)
    else:
        # 确保设置正确的大纲级别
        level = int(style_key[-1]) if style_key.startswith("heading") else 1
        ol = pPr.find(f"{{{W}}}outlineLvl")
        if ol is None:
            ol = OxmlElement("w:outlineLvl")
            ol.set(qn("w:val"), str(level - 1))  # 0-based
            pPr.append(ol)
        else:
            ol.set(qn("w:val"), str(level - 1))  # 0-based

    # run
    run = para.add_run(text)
    if hs is not None and hs.rPr is not None:
        # 复制 rPr 并确保字体颜色为黑色
        rPr_copy = copy.deepcopy(hs.rPr)
        # 移除现有的颜色设置
        color_elem = rPr_copy.find(f"{{{W}}}color")
        if color_elem is not None:
            rPr_copy.remove(color_elem)
        # 添加黑色字体设置
        color_elem = OxmlElement("w:color")
        color_elem.set(qn("w:val"), "000000")  # 黑色
        rPr_copy.append(color_elem)
        run._r.insert(0, rPr_copy)


def _apply_pPr(p_elem, pPr_proto) -> None:
    """完整替换段落 pPr（不继承）。sectPr 子元素保留（分节符不能丢）。"""
    if pPr_proto is None: return
    existing = p_elem.find(f"{{{W}}}pPr")
    # 保留 sectPr（分节符），其他全部替换
    saved_sectPr = None
    if existing is not None:
        saved_sectPr = existing.find(f"{{{W}}}sectPr")
        p_elem.remove(existing)
    new_pPr = copy.deepcopy(pPr_proto)
    # 去掉 pStyle（输出文档不继承样式）
    pStyle = new_pPr.find(f"{{{W}}}pStyle")
    if pStyle is not None:
        new_pPr.remove(pStyle)
    if saved_sectPr is not None:
        new_pPr.append(copy.deepcopy(saved_sectPr))
    p_elem.insert(0, new_pPr)
