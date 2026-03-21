from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import copy

from .constants import W


def _render_table(doc, item, st):
    caption   = item.get("caption", "")
    rows_data = item["data"]
    if not rows_data:
        return
    num_cols = max(len(r) for r in rows_data)

    if caption:
        # caption 格式完全来自模板 [[表格]] 块内 caption 行的 pPr/rPr
        # 约定：[[表格]] 块内第一行是 caption 示例行（在表格元素之前）
        cap = doc.add_paragraph()
        _apply_pPr(cap._p, st.table_caption_pPr)
        r = cap.add_run(caption)
        if st.table_caption_rPr is not None:
            r._r.insert(0, copy.deepcopy(st.table_caption_rPr))

    tbl = (_clone_table_with_data(st.table_proto, rows_data, num_cols)
           if st.table_proto is not None
           else _build_plain_table(rows_data, num_cols))

    body   = doc.element.body
    sectPr = body.find(f"{{{W}}}sectPr")
    if sectPr is not None:
        sectPr.addprevious(tbl)
    else:
        body.append(tbl)
    doc.add_paragraph("")


def _clone_table_with_data(proto, rows_data, num_cols):
    tbl = OxmlElement("w:tbl")
    proto_tblPr = proto.find(f"{{{W}}}tblPr")
    if proto_tblPr is not None:
        tbl.append(copy.deepcopy(proto_tblPr))
    proto_grid = proto.find(f"{{{W}}}tblGrid")
    if proto_grid is not None:
        grid = copy.deepcopy(proto_grid)
        cols = grid.findall(f"{{{W}}}gridCol")
        while len(cols) < num_cols:
            gc = OxmlElement("w:gridCol")
            gc.set(qn("w:w"), "1440")
            grid.append(gc)
            cols = grid.findall(f"{{{W}}}gridCol")
        tbl.append(grid)
    proto_rows   = proto.findall(f"{{{W}}}tr")
    header_proto = proto_rows[0] if proto_rows else None
    data_proto   = proto_rows[1] if len(proto_rows) > 1 else header_proto
    for row_idx, row_data in enumerate(rows_data):
        is_header = (row_idx == 0)
        tbl.append(_build_row_from_proto(
            header_proto if is_header else data_proto,
            row_data, num_cols, is_header))
    return tbl


def _build_row_from_proto(row_proto, row_data, num_cols, is_header):
    tr = OxmlElement("w:tr")
    if row_proto is not None:
        trPr = row_proto.find(f"{{{W}}}trPr")
        if trPr is not None:
            tr.append(copy.deepcopy(trPr))
    proto_cells = row_proto.findall(f"{{{W}}}tc") if row_proto is not None else []
    for col_idx in range(num_cols):
        cell_text  = row_data[col_idx] if col_idx < len(row_data) else ""
        cell_proto = (proto_cells[col_idx] if col_idx < len(proto_cells)
                      else (proto_cells[-1] if proto_cells else None))
        tr.append(_build_cell_from_proto(cell_proto, cell_text, is_header))
    return tr


def _build_cell_from_proto(cell_proto, text, is_header):
    tc = OxmlElement("w:tc")
    if cell_proto is not None:
        tcPr = cell_proto.find(f"{{{W}}}tcPr")
        if tcPr is not None:
            tc.append(copy.deepcopy(tcPr))
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    if cell_proto is not None:
        orig_p = cell_proto.find(f"{{{W}}}p")
        if orig_p is not None:
            pPr = orig_p.find(f"{{{W}}}pPr")
            if pPr is not None:
                p.append(copy.deepcopy(pPr))
            orig_r = orig_p.find(f"{{{W}}}r")
            if orig_r is not None:
                rPr = orig_r.find(f"{{{W}}}rPr")
                if rPr is not None:
                    new_rPr = copy.deepcopy(rPr)
                    if is_header and new_rPr.find(f"{{{W}}}b") is None:
                        new_rPr.insert(0, OxmlElement("w:b"))
                    r.append(new_rPr)
    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    r.append(t)
    p.append(r)
    tc.append(p)
    return tc


def _build_plain_table(rows_data, num_cols):
    total_w = 9360
    col_w   = total_w // num_cols
    col_ws  = [col_w] * num_cols
    col_ws[-1] += total_w - sum(col_ws)
    tbl = OxmlElement("w:tbl")
    tblPr = OxmlElement("w:tblPr")
    tblW  = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(total_w))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    tbl.append(tblPr)
    tblGrid = OxmlElement("w:tblGrid")
    for w in col_ws:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        tblGrid.append(gc)
    tbl.append(tblGrid)
    for row_idx, row_data in enumerate(rows_data):
        is_header = (row_idx == 0)
        tr = OxmlElement("w:tr")
        for col_idx in range(num_cols):
            cell_text = row_data[col_idx] if col_idx < len(row_data) else ""
            tc = OxmlElement("w:tc")
            tcPr = OxmlElement("w:tcPr")
            tcW  = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(col_ws[col_idx]))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)
            tc.append(tcPr)
            p = OxmlElement("w:p")
            r = OxmlElement("w:r")
            if is_header:
                rPr = OxmlElement("w:rPr")
                rPr.append(OxmlElement("w:b"))
                r.append(rPr)
            t = OxmlElement("w:t")
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t.text = cell_text
            r.append(t)
            p.append(r)
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)
    return tbl


def _apply_pPr(p_elem, pPr_proto) -> None:
    """完整替换段落 pPr（不继承）。sectPr 子元素保留（分节符不能丢）。"""
    if pPr_proto is None: return
    existing = p_elem.find(f"{{{W}}}pPr")
    # 保留 sectPr（分节符），其他全部替换
    saved_sectPr = None
    if existing is not None:
        saved_sectPr = existing.find(f"{{{W}}}sectPr")
        p_elem.remove(existing)
    new_pPr = copy.deepcopy(pPr_proto)
    # 去掉 pStyle（输出文档不继承样式）
    pStyle = new_pPr.find(f"{{{W}}}pStyle")
    if pStyle is not None:
        new_pPr.remove(pStyle)
    if saved_sectPr is not None:
        new_pPr.append(copy.deepcopy(saved_sectPr))
    p_elem.insert(0, new_pPr)
