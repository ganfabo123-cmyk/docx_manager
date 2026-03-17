from docx import Document
from pathlib import Path
import copy

from .toc_renderer import _render_toc_auto, _render_toc_manual, _add_heading_with_exclude
from .footer import apply_footer_config
from .section_renderer import _render_section
from .table_renderer import _render_table
from .formula_renderer import _render_formula
from .image_renderer import _render_image
from .references_renderer import _render_references
from .citation import insert_citation
from .section_break import _insert_section_break
from .constants import W
from docx.oxml.ns import qn
from ..models.models import StyleTemplate
from docx.oxml import OxmlElement


def generate(data: dict, st: StyleTemplate, output_path: str, footer_config: list[dict] | None = None) -> Document:
    """
    核心生成函数。

    文档装配顺序（固定）：
      前置章节（摘要等）
      [分节符，如果需要前/正文分页脚]
      目录（如果 content 里有 toc 块）
      正文内容
      后置章节（结论、致谢等）
      参考文献
    """
    toc_mode    = data.get("toc_mode", "auto")
    toc_entries = data.get("toc_entries", [])

    has_frontmatter = _has_frontmatter_sections(data.get("content", []))
    needs_split     = (has_frontmatter and footer_config is not None and
                       any(c["section"] == "frontmatter" for c in footer_config))

    doc = Document()

    split_inserted = False
    for item in data.get("content", []):
        t = item["type"]

        # 在第一个正文 heading1 前插入分节符
        if needs_split and not split_inserted and t == "heading1":
            _insert_section_break(doc)
            split_inserted = True

        if t in ("heading1", "heading2", "heading3"):
            exclude = item.get("toc_exclude", False)
            _add_heading_with_exclude(doc, item["value"], t, exclude=exclude, st=st)

        elif t == "body":
            p = doc.add_paragraph()
            _apply_pPr(p._p, st.body_pPr)
            run = p.add_run(item["value"])
            if st.body_rPr is not None:
                # 复制 rPr 并确保字体颜色为黑色
                rPr_copy = copy.deepcopy(st.body_rPr)
                # 移除现有的颜色设置
                color_elem = rPr_copy.find(f"{{{W}}}color")
                if color_elem is not None:
                    rPr_copy.remove(color_elem)
                # 添加黑色字体设置
                color_elem = OxmlElement("w:color")
                color_elem.set(qn("w:val"), "000000")  # 黑色
                rPr_copy.append(color_elem)
                run._r.insert(0, rPr_copy)

        elif t == "table":
            _render_table(doc, item, st)

        elif t == "formula":
            _render_formula(doc, item, st)

        elif t == "image":
            _render_image(doc, item, st)

        elif t == "section":
            _render_section(doc, item, st)

        elif t == "toc":
            if toc_mode == "manual" and toc_entries:
                _render_toc_manual(doc, item, st, toc_entries)
            else:
                _render_toc_auto(doc, item, st)

    # 参考文献
    refs = data.get("references", [])
    if refs:
        _render_references(doc, refs, st)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    # 引用上标
    doc = Document(output_path)
    for cite in data.get("citations", []):
        insert_citation(doc, cite["ref_id"], (cite["before"], cite.get("after", "")))

    # 页脚
    if footer_config:
        apply_footer_config(doc, footer_config, has_frontmatter and split_inserted)

    doc.save(output_path)
    return doc


def _has_frontmatter_sections(content: list[dict]) -> bool:
    for item in content:
        if item["type"] == "heading1":
            return False
        if item["type"] == "section":
            return True
    return False


def _apply_pPr(p_elem, pPr_proto) -> None:
    """完整替换段落 pPr（不继承）。sectPr 子元素保留（分节符不能丢）。"""
    if pPr_proto is None: return
    existing = p_elem.find(f"{{{W}}}pPr")
    # 保留 sectPr（分节符），其他全部替换
    saved_sectPr = None
    if existing is not None:
        saved_sectPr = existing.find(f"{{{W}}}sectPr")
        p_elem.remove(existing)
    new_pPr = copy.deepcopy(pPr_proto)
    # 去掉 pStyle（输出文档不继承样式）
    pStyle = new_pPr.find(f"{{{W}}}pStyle")
    if pStyle is not None:
        new_pPr.remove(pStyle)
    if saved_sectPr is not None:
        new_pPr.append(copy.deepcopy(saved_sectPr))
    p_elem.insert(0, new_pPr)
