import json
import traceback
from typing import List, Dict, Any

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx not installed, docx generation will be disabled")

from docx import Document


# 定义必要的命名空间
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M = "http://schemas.openxmlformats.org/officeDocument/2006/math" # 公式命名空间


def identify_short_blocks(blocks: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Identify short text blocks
    1. First pass: collect blocks with content length < 30 (A)
    2. Second pass: collect blocks that have corresponding TOC items (B)
    3. If B exists, return B; otherwise return A
    """
    try:
        # First pass: collect all blocks with content length < 30 (A)
        A = [block for block in blocks if len(block['content']) < 30]
        
        # Second pass: collect blocks that have corresponding TOC items (B)
        B = []
        # Build content map: base_content -> list of blocks
        content_map = {}
        for block in blocks:
            # Remove trailing \t and numbers to get base content
            base_content = block['content'].split('\t')[0].strip()
            if base_content not in content_map:
                content_map[base_content] = []
            content_map[base_content].append(block)
        
        # Add blocks that have corresponding TOC items
        for base_content, block_list in content_map.items():
            if len(block_list) > 1:
                # Check if there's at least one TOC item (ends with \t[0-9]+)
                has_toc_item = any(
                    ('\t' in block['content'] and block['content'].split('\t')[-1].strip() != "") or 
                    ("".join(block['content'].split()) == "目录") 
                    for block in block_list
                )                

                if has_toc_item:
                    for block in block_list:
                        # Add the non-TOC item (heading) to candidates
                        if '\t' not in block['content']:
                            B.append(block)
        
        # If B exists, return B; otherwise return A
        if B:
            return B
        else:
            return A
    except Exception as e:
        print(f"Error identifying short blocks: {e}")
        print(traceback.format_exc())
        return []

def save_parsed_styles(analysis: List[Dict[str, Any]], output_path: str) -> None:
    """
    Save parsed styles to JSON file
    """
    try:
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(analysis, f, ensure_ascii=False, indent=2)
        print(f"Parsed styles saved to {output_path}")
    except Exception as e:
        print(f"Error saving parsed styles: {e}")
        print(traceback.format_exc())

def generate_docx_document(parsed_styles: List[Dict[str, Any]], output_path: str) -> bool:
    """
    Generate docx document from parsed styles
    """
    try:
        doc = Document()
        
        for item in parsed_styles:
            if item['type'] == 'h1':
                # Level 1 heading
                para = doc.add_heading(item['content'], level=1)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif item['type'] == 'h2':
                # Level 2 heading
                para = doc.add_heading(item['content'], level=2)
            elif item['type'] == 'h3':
                # Level 3 heading
                para = doc.add_heading(item['content'], level=3)
            elif item['type'] in ['toc1', 'toc2', 'toc3']:
                # Table of contents item
                para = doc.add_paragraph(item['content'])
            else:
                # Regular paragraph
                para = doc.add_paragraph(item['content'])
        
        doc.save(output_path)
        print(f"Generated docx document saved to: {output_path}")
        return True
    except Exception as e:
        print(f"Error generating docx document: {e}")
        print(traceback.format_exc())
        return False