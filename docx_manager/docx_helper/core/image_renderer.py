from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from pathlib import Path
import io
import base64
import copy

from .constants import _ALIGN_MAP, W


def _render_image(doc, item, st):
    width_in = float(item.get("width", 4.0))
    align    = _ALIGN_MAP.get(item.get("align", "center"), WD_ALIGN_PARAGRAPH.CENTER)
    caption  = item.get("caption", "")
    img_stream = _get_image_stream(item)
    if img_stream is None:
        ph = doc.add_paragraph()
        _apply_pPr(ph._p, st.image_pPr)
        ph.add_run(f"[图片占位：{caption or item.get('path','')}]")
        return
    para = doc.add_paragraph()
    para.alignment = align
    _apply_pPr(para._p, st.image_pPr)
    run = para.add_run()
    try:
        run.add_picture(img_stream, width=Inches(width_in))
    except Exception as exc:
        para.clear()
        para.add_run(f"[图片插入失败：{exc}]")
    if caption:
        cap_para = doc.add_paragraph()
        # caption 行格式来自模板 [[图片]] 块第二行（image_caption_pPr）
        _apply_pPr(cap_para._p, st.image_caption_pPr if st.image_caption_pPr is not None else st.image_pPr)
        r = cap_para.add_run(caption)
        if st.caption_rPr is not None:
            r._r.insert(0, copy.deepcopy(st.caption_rPr))


def _get_image_stream(item):
    if "path" in item:
        p = Path(item["path"])
        return io.BytesIO(p.read_bytes()) if p.exists() else None
    if "base64" in item:
        try:
            return io.BytesIO(base64.b64decode(item["base64"]))
        except Exception:
            return None
    return None


def _apply_pPr(p_elem, pPr_proto) -> None:
    """完整替换段落 pPr（不继承）。sectPr 子元素保留（分节符不能丢）。"""
    if pPr_proto is None: return
    existing = p_elem.find(f"{{{W}}}pPr")
    # 保留 sectPr（分节符），其他全部替换
    saved_sectPr = None
    if existing is not None:
        saved_sectPr = existing.find(f"{{{W}}}sectPr")
        p_elem.remove(existing)
    new_pPr = copy.deepcopy(pPr_proto)
    # 去掉 pStyle（输出文档不继承样式）
    pStyle = new_pPr.find(f"{{{W}}}pStyle")
    if pStyle is not None:
        new_pPr.remove(pStyle)
    if saved_sectPr is not None:
        new_pPr.append(copy.deepcopy(saved_sectPr))
    p_elem.insert(0, new_pPr)
