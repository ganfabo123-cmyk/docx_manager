import sys
import os
from pathlib import Path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'docx_helper'))

try:
    from docx import Document
    print("✓ python-docx 模块导入成功")
    
    data_dir = Path(__file__).parent / "data"
    doc_path = data_dir / "template.docx"
    
    print(f"正在解析文档: {doc_path}")
    
    doc = Document(str(doc_path))
    print(f"✓ 文档加载成功，共 {len(doc.paragraphs)} 个段落，{len(doc.tables)} 个表格")
    
    from docx.oxml.ns import qn
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    
    blocks = []
    block_id = 0
    
    for element in doc.element.body:
        block_id += 1
        
        if element.tag == f"{{{W}}}p":
            from docx.text.paragraph import Paragraph
            paragraph = Paragraph(element, doc)
            
            if not paragraph.text.strip():
                continue
            
            style_info = {
                "block_id": block_id,
                "type": "paragraph",
                "text": paragraph.text.strip()[:50] + "..." if len(paragraph.text.strip()) > 50 else paragraph.text.strip(),
                "style_name": paragraph.style.name if paragraph.style else "Normal",
            }
            
            # 提取段落级别的格式 (pPr)
            p_elem = paragraph._element
            pPr = p_elem.find(f"{{{W}}}pPr")
            
            if pPr is not None:
                # 对齐方式
                jc = pPr.find(f"{{{W}}}jc")
                if jc is not None:
                    jc_val = jc.get(qn("w:val"))
                    align_map = {"both": "both", "center": "center", "distribute": "distribute", 
                               "end": "right", "highKashida": "highKashida", 
                               "left": "left", "lowKashida": "lowKashida", 
                               "mediumKashida": "mediumKashida", "numTab": "numTab", 
                               "right": "right", "start": "left", "thaiDistribute": "thaiDistribute"}
                    style_info["alignment"] = align_map.get(jc_val, jc_val)
                
                # 行间距
                spacing = pPr.find(f"{{{W}}}spacing")
                if spacing is not None:
                    style_info["spacing"] = {}
                    if spacing.get(qn("w:line")):
                        style_info["spacing"]["line"] = spacing.get(qn("w:line"))
                    if spacing.get(qn("w:before")):
                        style_info["spacing"]["before"] = spacing.get(qn("w:before"))
                    if spacing.get(qn("w:after")):
                        style_info["spacing"]["after"] = spacing.get(qn("w:after"))
                
                # 缩进
                ind = pPr.find(f"{{{W}}}ind")
                if ind is not None:
                    style_info["indentation"] = {}
                    if ind.get(qn("w:firstLine")):
                        style_info["indentation"]["firstLine"] = ind.get(qn("w:firstLine"))
                    if ind.get(qn("w:left")):
                        style_info["indentation"]["left"] = ind.get(qn("w:left"))
            
            # 提取字符级别的格式 (rPr)
            if paragraph.runs:
                run = paragraph.runs[0]
                rPr = run._r
                
                if rPr is not None:
                    # 字体大小
                    sz = rPr.find(f"{{{W}}}sz")
                    if sz is not None:
                        sz_val = sz.get(qn("w:val"))
                        if sz_val:
                            style_info["font_size"] = f"{int(sz_val) / 2}pt"
                    
                    # 字体名称
                    rFonts = rPr.find(f"{{{W}}}rFonts")
                    if rFonts is not None:
                        ascii_font = rFonts.get(qn("w:ascii"))
                        hAnsi_font = rFonts.get(qn("w:hAnsi"))
                        style_info["font_name"] = ascii_font or hAnsi_font or "Default"
                    
                    # 加粗
                    b = rPr.find(f"{{{W}}}b")
                    if b is not None:
                        style_info["bold"] = b.get(qn("w:val")) != "0"
                    
                    # 斜体
                    i = rPr.find(f"{{{W}}}i")
                    if i is not None:
                        style_info["italic"] = i.get(qn("w:val")) != "0"
                    
                    # 字体颜色
                    color = rPr.find(f"{{{W}}}color")
                    if color is not None:
                        style_info["color"] = color.get(qn("w:val"))
            
            blocks.append(style_info)
        
        elif element.tag == f"{{{W}}}tbl":
            from docx.table import Table
            table = Table(element, doc.element.body)
            
            style_info = {
                "block_id": block_id,
                "type": "table",
                "rows": len(table.rows),
                "cols": len(table.columns) if table.rows else 0,
            }
            
            # 查找表格前的段落（可能是标题）
            tbl_elem = table._element
            prev_elem = tbl_elem.getprevious()
            if prev_elem is not None and prev_elem.tag == f"{{{W}}}p":
                prev_para = Paragraph(prev_elem, doc)
                style_info["caption"] = prev_para.text.strip()
            
            # 表格样式
            tblPr = tbl_elem.find(f"{{{W}}}tblPr")
            if tblPr is not None:
                style_info["table_style"] = {}
                
                # 表格宽度
                tblW = tblPr.find(f"{{{W}}}tblW")
                if tblW is not None:
                    style_info["table_style"]["width"] = tblW.get(qn("w:w"))
                    style_info["table_style"]["width_type"] = tblW.get(qn("w:type"))
                
                # 边框
                tblBorders = tblPr.find(f"{{{W}}}tblBorders")
                if tblBorders is not None:
                    style_info["table_style"]["borders"] = {}
                    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                        border = tblBorders.find(f"{{{W}}}{border_name}")
                        if border is not None:
                            style_info["table_style"]["borders"][border_name] = {
                                "val": border.get(qn("w:val")),
                                "sz": border.get(qn("w:sz")),
                                "color": border.get(qn("w:color"))
                            }
            
            blocks.append(style_info)
    
    print(f"\n共找到 {len(blocks)} 个块\n")
    print("=" * 80)
    
    # 显示前 20 个块的详细信息
    for i, block in enumerate(blocks[:20], 1):
        print(f"\n块 #{i} (ID: {block['block_id']})")
        print(f"类型: {block['type']}")
        
        if block['type'] == 'paragraph':
            print(f"文本: {block['text']}")
            print(f"样式名称: {block['style_name']}")
            
            if 'alignment' in block:
                print(f"对齐: {block['alignment']}")
            if 'font_size' in block:
                print(f"字号: {block['font_size']}")
            if 'font_name' in block:
                print(f"字体: {block['font_name']}")
            if 'bold' in block:
                print(f"加粗: {block['bold']}")
            if 'italic' in block:
                print(f"斜体: {block['italic']}")
            if 'color' in block:
                print(f"颜色: #{block['color']}")
            if 'spacing' in block:
                print(f"间距: {block['spacing']}")
            if 'indentation' in block:
                print(f"缩进: {block['indentation']}")
        
        elif block['type'] == 'table':
            print(f"表格: {block['rows']} 行 x {block['cols']} 列")
            if 'caption' in block:
                print(f"标题: {block['caption']}")
            if 'table_style' in block:
                print(f"表格样式: {block['table_style']}")
    
    if len(blocks) > 20:
        print(f"\n... 还有 {len(blocks) - 20} 个块未显示")
    
    # 保存为 JSON
    import json
    output_path = data_dir / "parsed_styles.json"
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(blocks, f, ensure_ascii=False, indent=2)
    
    print(f"\n{'=' * 80}")
    print(f"\n解析完成！结果已保存到: {output_path}")
    
except ImportError as e:
    print(f"✗ 导入错误: {e}")
    print("请确保 docx_helper 目录存在且包含 python-docx 模块")
except Exception as e:
    print(f"✗ 错误: {e}")
    import traceback
    traceback.print_exc()
