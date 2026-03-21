from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
import copy

from .constants import W, SECTION_DEFAULT_TITLES
from .toc_renderer import _add_heading_with_exclude
from ..models.models import StyleTemplate


def _render_section(doc: Document, item: dict, st: StyleTemplate) -> None:
    stype   = item.get("section_type", "custom")
    title   = item.get("title") or SECTION_DEFAULT_TITLES.get(stype, stype)
    value   = item.get("value", "")
    ss      = st.section_styles.get(stype, {})
    exclude = item.get("toc_exclude", False)

    # 标题
    title_para = doc.add_paragraph()
    title_pPr  = ss.get("title_pPr")
    if title_pPr is not None:
        # 应用章节标题样式，并添加适当的标题样式
        _apply_pPr_with_style(title_para._p, title_pPr, "Heading1")
        if exclude:
            pPr = title_para._p.find(f"{{{W}}}pPr")
            if pPr is not None:
                ol = OxmlElement("w:outlineLvl")
                ol.set(qn("w:val"), "9")
                pPr.append(ol)
        r = title_para.add_run(title)
        title_rPr = ss.get("title_rPr")
        if title_rPr is not None:
            # 复制 rPr 并确保字体颜色为黑色
            rPr_copy = copy.deepcopy(title_rPr)
            # 移除现有的颜色设置
            color_elem = rPr_copy.find(f"{{{W}}}color")
            if color_elem is not None:
                rPr_copy.remove(color_elem)
            # 添加黑色字体设置
            color_elem = OxmlElement("w:color")
            color_elem.set(qn("w:val"), "000000")  # 黑色
            rPr_copy.append(color_elem)
            r._r.insert(0, rPr_copy)
    else:
        # 无独立章节标题样式 → 复用模板 [[一级标题]] 样式
        _add_heading_with_exclude(doc, title, "heading1", exclude=exclude, st=st)

    # 正文：优先章节自己的 body_pPr/rPr，无则不写格式（保持 Word 默认）
    # 注意：不回退到全局 body 样式，否则正文格式会污染章节正文
    body_pPr = ss.get("body_pPr")
    body_rPr = ss.get("body_rPr")
    for para_text in _split_paragraphs(value):
        p = doc.add_paragraph()
        _apply_pPr(p._p, body_pPr)
        run = p.add_run(para_text)
        if body_rPr is not None:
            # 复制 rPr 并确保字体颜色为黑色
            rPr_copy = copy.deepcopy(body_rPr)
            # 移除现有的颜色设置
            color_elem = rPr_copy.find(f"{{{W}}}color")
            if color_elem is not None:
                rPr_copy.remove(color_elem)
            # 添加黑色字体设置
            color_elem = OxmlElement("w:color")
            color_elem.set(qn("w:val"), "000000")  # 黑色
            rPr_copy.append(color_elem)
            run._r.insert(0, rPr_copy)


def _apply_pPr_with_style(p_elem, pPr_proto, style_id) -> None:
    """应用段落样式，同时保留指定的样式ID。"""
    if pPr_proto is None:
        return
    existing = p_elem.find(f"{{{W}}}pPr")
    # 保留 sectPr（分节符），其他全部替换
    saved_sectPr = None
    if existing is not None:
        saved_sectPr = existing.find(f"{{{W}}}sectPr")
        p_elem.remove(existing)
    new_pPr = copy.deepcopy(pPr_proto)
    # 确保添加 pStyle
    pStyle = new_pPr.find(f"{{{W}}}pStyle")
    if pStyle is None:
        pStyle = OxmlElement("w:pStyle")
        pStyle.set(qn("w:val"), style_id)
        new_pPr.insert(0, pStyle)
    else:
        pStyle.set(qn("w:val"), style_id)
    if saved_sectPr is not None:
        new_pPr.append(copy.deepcopy(saved_sectPr))
    p_elem.insert(0, new_pPr)


def _split_paragraphs(text: str) -> list[str]:
    paras = re.split(r"\n{2,}", text)
    return [p.strip().replace("\n", " ") for p in paras if p.strip()]


def _apply_pPr(p_elem, pPr_proto) -> None:
    """完整替换段落 pPr（不继承）。sectPr 子元素保留（分节符不能丢）。"""
    if pPr_proto is None: return
    existing = p_elem.find(f"{{{W}}}pPr")
    # 保留 sectPr（分节符），其他全部替换
    saved_sectPr = None
    if existing is not None:
        saved_sectPr = existing.find(f"{{{W}}}sectPr")
        p_elem.remove(existing)
    new_pPr = copy.deepcopy(pPr_proto)
    # 去掉 pStyle（输出文档不继承样式）
    pStyle = new_pPr.find(f"{{{W}}}pStyle")
    if pStyle is not None:
        new_pPr.remove(pStyle)
    if saved_sectPr is not None:
        new_pPr.append(copy.deepcopy(saved_sectPr))
    p_elem.insert(0, new_pPr)
