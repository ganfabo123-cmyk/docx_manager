from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import copy

from .constants import W
from .toc_renderer import _add_heading_with_exclude as _add_heading


def _render_references(doc, refs, st):
    _add_heading(doc, "参考文献", "heading1", st=st)
    for ref in refs:
        para = doc.add_paragraph()
        pPr  = para._p.get_or_add_pPr()
        _apply_pPr(para._p, st.ref_pPr_proto)
        run_num = para.add_run(f"[{ref['id']}] ")
        if st.ref_rPr_proto is not None:
            run_num._r.insert(0, copy.deepcopy(st.ref_rPr_proto))
        run_num.bold = True
        r_text = para.add_run(ref["text"])
        if st.ref_rPr_proto is not None:
            r_text._r.insert(0, copy.deepcopy(st.ref_rPr_proto))


def _apply_pPr(p_elem, pPr_proto) -> None:
    """完整替换段落 pPr（不继承）。sectPr 子元素保留（分节符不能丢）。"""
    if pPr_proto is None: return
    existing = p_elem.find(f"{{{W}}}pPr")
    # 保留 sectPr（分节符），其他全部替换
    saved_sectPr = None
    if existing is not None:
        saved_sectPr = existing.find(f"{{{W}}}sectPr")
        p_elem.remove(existing)
    new_pPr = copy.deepcopy(pPr_proto)
    # 去掉 pStyle（输出文档不继承样式）
    pStyle = new_pPr.find(f"{{{W}}}pStyle")
    if pStyle is not None:
        new_pPr.remove(pStyle)
    if saved_sectPr is not None:
        new_pPr.append(copy.deepcopy(saved_sectPr))
    p_elem.insert(0, new_pPr)
