import re
import json
import urllib.request

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx.enum.text import WD_ALIGN_PARAGRAPH
from .constants import W, _ALIGN_MAP, FOOTER_STYLES, _FOOTER_LLM_SYSTEM


def parse_footer_spec(spec: str, api_key: str | None = None) -> list[dict]:
    if api_key:
        return _parse_footer_via_llm(spec, api_key)
    return _parse_footer_heuristic(spec)


def _parse_footer_via_llm(spec: str, api_key: str) -> list[dict]:
    payload = json.dumps({
        "model": "claude-sonnet-4-20250514",
        "max_tokens": 512,
        "system": _FOOTER_LLM_SYSTEM,
        "messages": [{"role": "user", "content": spec}],
    }).encode()
    req = urllib.request.Request(
        "https://api.anthropic.com/v1/messages",
        data=payload,
        headers={
            "Content-Type":      "application/json",
            "x-api-key":         api_key,
            "anthropic-version": "2023-06-01",
        },
        method="POST",
    )
    try:
        with urllib.request.urlopen(req, timeout=15) as resp:
            body = json.loads(resp.read())
        text = body["content"][0]["text"].strip()
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text)
        config = json.loads(text)
        validated = []
        for it in config:
            if it.get("style") in FOOTER_STYLES and \
               it.get("section") in ("frontmatter", "mainmatter"):
                validated.append({
                    "section": it["section"],
                    "style":   it["style"],
                    "start":   int(it.get("start", 1)),
                })
        return validated
    except Exception as exc:
        print(f"⚠️  LLM 页脚解析失败（{exc}），退回规则匹配")
        return _parse_footer_heuristic(spec)


def _parse_footer_heuristic(spec: str) -> list[dict]:
    spec_l = spec.lower()

    def _detect_style(text: str) -> str:
        if "罗马" in text or "roman" in text.lower():
            return "roman_upper_center" if ("大写" in text or "upper" in text.lower()) \
                   else "roman_lower_center"
        if "-" in text and ("x" in text.lower() or "页" in text):
            return "arabic_dash"
        if "第" in text and "页" in text:
            return "arabic_page_x"
        if "/" in text:
            return "arabic_slash"
        if "无" in text or "none" in text.lower():
            return "none"
        return "arabic_center"

    front_kws = ["前置", "封面", "目录", "摘要", "前言", "frontmatter"]
    main_kws  = ["正文", "章节", "mainmatter", "阿拉伯"]
    result    = []

    if any(kw in spec_l for kw in front_kws):
        for kw in front_kws:
            idx = spec_l.find(kw)
            if idx != -1:
                snippet = spec[max(0, idx-5):idx+40]
                result.append({"section": "frontmatter",
                                "style": _detect_style(snippet), "start": 1})
                break

    if any(kw in spec_l for kw in main_kws):
        for kw in main_kws:
            idx = spec_l.find(kw)
            if idx != -1:
                snippet = spec[max(0, idx-5):idx+40]
                result.append({"section": "mainmatter",
                                "style": _detect_style(snippet), "start": 1})
                break

    if not result:
        result.append({"section": "mainmatter",
                        "style": _detect_style(spec), "start": 1})
    return result


def apply_footer_config(doc: Document, config: list[dict], has_frontmatter: bool) -> None:
    front_cfg = next((c for c in config if c["section"] == "frontmatter"), None)
    main_cfg  = next((c for c in config if c["section"] == "mainmatter"),  None)
    sections  = doc.sections

    if has_frontmatter and len(sections) >= 2:
        if front_cfg:
            _write_section_footer(sections[0], front_cfg["style"],
                                  front_cfg.get("start", 1))
        if main_cfg:
            _write_section_footer(sections[1], main_cfg["style"],
                                  main_cfg.get("start", 1),
                                  unlink_from_prev=True)
    else:
        cfg = main_cfg or front_cfg
        if cfg:
            _write_section_footer(sections[-1], cfg["style"],
                                  cfg.get("start", 1))


def _write_section_footer(section, style_name: str, start_num: int = 1, unlink_from_prev: bool = False) -> None:
    style  = FOOTER_STYLES.get(style_name)
    sectPr = section._sectPr
    for old in sectPr.findall(f"{{{W}}}pgNumType"):
        sectPr.remove(old)

    if style is None:
        footer = section.footer
        footer.is_linked_to_previous = False
        for para in footer.paragraphs:
            para.clear()
        return

    pgNumType = OxmlElement("w:pgNumType")
    pgNumType.set(qn("w:fmt"),   style["fmt"])
    pgNumType.set(qn("w:start"), str(start_num))
    sectPr.append(pgNumType)

    footer = section.footer
    if unlink_from_prev:
        footer.is_linked_to_previous = False
    for para in footer.paragraphs:
        para.clear()
    fp = footer.paragraphs[0]
    fp.alignment = _ALIGN_MAP.get(style.get("align", "center"),
                                  WD_ALIGN_PARAGRAPH.CENTER)

    if style.get("prefix"):
        fp.add_run(style["prefix"])
    _add_field_run(fp, style.get("instr", "PAGE"))
    if style.get("with_numpages"):
        fp.add_run("/")
        _add_field_run(fp, "NUMPAGES")
    if style.get("suffix"):
        fp.add_run(style["suffix"])


def _add_field_run(para, instr: str) -> None:
    r1 = para.add_run()
    fc1 = OxmlElement("w:fldChar")
    fc1.set(qn("w:fldCharType"), "begin")
    r1._r.append(fc1)

    r2 = para.add_run()
    it = OxmlElement("w:instrText")
    it.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    it.text = f" {instr} "
    r2._r.append(it)

    r3 = para.add_run()
    fc2 = OxmlElement("w:fldChar")
    fc2.set(qn("w:fldCharType"), "end")
    r3._r.append(fc2)
