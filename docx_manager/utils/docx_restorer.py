"""
DOCX文档还原器
从JSON格式还原为docx文档
"""
import json
import base64
import io
import traceback
import copy
from typing import Dict, List, Any, Optional
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Twips, Cm

_DEFAULT_IMG_WIDTH_CM = 12.00
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.part import Part
from docx.opc.packuri import PackURI
from lxml import etree


W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
V_NS = "urn:schemas-microsoft-com:vml"
O_NS = "urn:schemas-microsoft-com:office:office"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


class DocxRestorer:
    """DOCX文档还原器"""
    
    HEADING_STYLE_MAP = {
        1: "Heading 1",
        2: "Heading 2", 
        3: "Heading 3"
    }
    
    ALIGN_MAP = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
    }
    
    def __init__(self, elements: List[Dict[str, Any]], citations: List[Dict[str, Any]] = None):
        self.elements = elements
        self.citations = citations or []
        self.doc = Document()
        self._ole_counter = 0
        self._shape_counter = 0
    
    def load_citations(self, citations_path: str) -> None:
        if Path(citations_path).exists():
            with open(citations_path, 'r', encoding='utf-8') as f:
                self.citations = json.load(f)
    
    def _restore_citations(self):
        for citation in self.citations:
            ref_id = citation.get("ref_id")
            before = citation.get("before", "")
            after = citation.get("after", "")
            element_id = citation.get("element_id")
            
            self._insert_superscript_citation(ref_id, (before, after))
    
    def _insert_superscript_citation(self, ref_id: int, context: tuple) -> bool:
        before, after = context
        old_ref_text = f"[{ref_id}]"
        
        for para in self.doc.paragraphs:
            full = para.text
            pos = full.find(before)
            if pos == -1:
                continue
            
            insert_at = pos + len(before)
            
            if after and after[:5] not in full[insert_at:]:
                continue
            
            runs = para.runs
            if not runs:
                continue
            
            cur = 0
            target_idx, target_off = len(runs) - 1, len(runs[-1].text)
            for ri, run in enumerate(runs):
                end = cur + len(run.text)
                if cur <= insert_at <= end:
                    target_idx = ri
                    target_off = insert_at - cur
                    break
                cur = end
            
            target_run = runs[target_idx]
            orig_text = target_run.text
            target_run.text = orig_text[:target_off]
            
            orig_rPr = target_run._r.find(qn("w:rPr"))
            
            r_sup = OxmlElement("w:r")
            new_rPr = copy.deepcopy(orig_rPr) if orig_rPr is not None else OxmlElement("w:rPr")
            
            va = OxmlElement("w:vertAlign")
            va.set(qn("w:val"), "superscript")
            new_rPr.append(va)
            r_sup.append(new_rPr)
            
            t_sup = OxmlElement("w:t")
            t_sup.text = old_ref_text
            r_sup.append(t_sup)
            
            r_tail = OxmlElement("w:r")
            if orig_rPr is not None:
                r_tail.append(copy.deepcopy(orig_rPr))
            
            t_tail = OxmlElement("w:t")
            t_tail.set(qn("xml:space"), "preserve")
            
            remaining_text = orig_text[target_off:]
            if remaining_text.startswith(old_ref_text):
                t_tail.text = remaining_text[len(old_ref_text):]
            else:
                t_tail.text = remaining_text
            
            r_tail.append(t_tail)
            
            target_run._r.addnext(r_tail)
            target_run._r.addnext(r_sup)
            
            return True
        
        return False
    
    def _apply_paragraph_style(self, paragraph, style: Dict[str, Any]):
        style_name = style.get("style_name", "Normal")
        
        if style.get("is_heading") and style.get("heading_level"):
            level = style["heading_level"]
            heading_style = self.HEADING_STYLE_MAP.get(level)
            if heading_style:
                try:
                    paragraph.style = heading_style
                except Exception:
                    pass
        
        alignment = style.get("alignment")
        if alignment and alignment in self.ALIGN_MAP:
            paragraph.alignment = self.ALIGN_MAP[alignment]
        
        pPr = paragraph._p.get_or_add_pPr()
        
        spacing_before = style.get("spacing_before")
        spacing_after = style.get("spacing_after")
        line_spacing = style.get("line_spacing")
        
        if any([spacing_before, spacing_after, line_spacing]):
            spacing = pPr.find(f"{{{W}}}spacing")
            if spacing is None:
                spacing = OxmlElement("w:spacing")
                pPr.append(spacing)
            
            if spacing_before:
                spacing.set(qn("w:before"), str(spacing_before))
            if spacing_after:
                spacing.set(qn("w:after"), str(spacing_after))
            if line_spacing:
                spacing.set(qn("w:line"), str(line_spacing))
        
        first_line_indent = style.get("first_line_indent")
        left_indent = style.get("left_indent")
        
        if any([first_line_indent, left_indent]):
            ind = pPr.find(f"{{{W}}}ind")
            if ind is None:
                ind = OxmlElement("w:ind")
                pPr.append(ind)
            
            if first_line_indent:
                ind.set(qn("w:firstLine"), str(first_line_indent))
            if left_indent:
                ind.set(qn("w:left"), str(left_indent))
    
    def _restore_paragraph(self, element: Dict[str, Any]):
        content = element.get("content", "")
        style = element.get("style", {})
        elem_type = element.get("type", "body")
        
        paragraph = self.doc.add_paragraph()
        
        if elem_type.startswith("heading"):
            try:
                level = int(elem_type.replace("heading", ""))
                heading_style = self.HEADING_STYLE_MAP.get(level)
                if heading_style:
                    paragraph.style = heading_style
            except (ValueError, KeyError):
                pass
        
        self._apply_paragraph_style(paragraph, style)
        
        paragraph.add_run(content)
    
    def _restore_table(self, element: Dict[str, Any]):
        content = element.get("content", [])
        style = element.get("style", {})
        
        if not content:
            return
        
        rows = len(content)
        cols = max(len(r) for r in content) if content else 0
        
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"
        
        for row_idx, row_data in enumerate(content):
            for col_idx, cell_content in enumerate(row_data):
                if col_idx < cols:
                    cell = table.rows[row_idx].cells[col_idx]
                    if isinstance(cell_content, dict) and cell_content.get('omath'):
                        paragraph = cell.paragraphs[0]
                        paragraph.clear()
                        if cell_content.get('text_before'):
                            paragraph.add_run(cell_content['text_before'])
                        try:
                            omath_elem = etree.fromstring(cell_content['omath'].encode())
                            paragraph._p.append(omath_elem)
                        except Exception:
                            paragraph.add_run('[公式]')
                        if cell_content.get('text_after'):
                            paragraph.add_run(cell_content['text_after'])
                    else:
                        cell.text = str(cell_content)
    
    def _restore_image(self, element: Dict[str, Any]):
        base64_str = element.get("base64", "")
        caption = element.get("caption", "")
        width_inches = element.get("width")

        if not base64_str:
            paragraph = self.doc.add_paragraph()
            paragraph.add_run("[图片占位符]")
            return

        try:
            from PIL import Image as PILImage
            image_bytes = base64.b64decode(base64_str)
            pil_img = PILImage.open(io.BytesIO(image_bytes))
            image_stream = io.BytesIO()
            pil_img.save(image_stream, format='PNG')
            image_stream.seek(0)

            paragraph = self.doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = paragraph.add_run()

            if width_inches:
                run.add_picture(image_stream, width=Inches(width_inches))
            else:
                run.add_picture(image_stream, width=Cm(_DEFAULT_IMG_WIDTH_CM))
            
            if caption:
                caption_para = self.doc.add_paragraph()
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_para.add_run(caption)
                
        except Exception as e:
            print(f"还原图片时出错: {e}")
            traceback.print_exc()
            paragraph = self.doc.add_paragraph()
            paragraph.add_run(f"[图片还原失败: {e}]")
    
    def _restore_formula(self, element: Dict[str, Any]):
        ole_base64 = element.get("ole_base64", "")
        
        if ole_base64:
            self._restore_ole_formula(element)
            return
        
        formula = element.get("formula", {})
        omml_str = formula.get("omath", "") or element.get("omml", "")
        label = formula.get("label", "") or element.get("label", "")
        is_inline = element.get("type") == "formula_inline"

        if is_inline:
            try:
                paragraph = self.doc.add_paragraph()
                self._apply_paragraph_style(paragraph, element.get('style', {}))
                text_before = formula.get('text_before', '')
                text_after = formula.get('text_after', '')
                if text_before:
                    paragraph.add_run(text_before)
                if omml_str:
                    omml_elem = etree.fromstring(omml_str.encode())
                    paragraph._p.append(omml_elem)
                if text_after:
                    paragraph.add_run(text_after)
            except Exception as e:
                print(f"还原行内公式时出错: {e}")
                traceback.print_exc()
                paragraph = self.doc.add_paragraph()
                paragraph.add_run(element.get('content', '[行内公式还原失败]'))
            return

        if omml_str:
            try:
                omml_elem = etree.fromstring(omml_str.encode())

                p = OxmlElement("w:p")
                pPr = OxmlElement("w:pPr")
                p.append(pPr)

                jc = OxmlElement("w:jc")
                jc.set(qn("w:val"), "center")
                pPr.append(jc)

                p.append(omml_elem)

                if label:
                    run = OxmlElement("w:r")
                    t = OxmlElement("w:t")
                    t.text = f"    {label}"
                    run.append(t)
                    p.append(run)

                body = self.doc.element.body
                sectPr = body.find(f"{{{W}}}sectPr")
                if sectPr is not None:
                    sectPr.addprevious(p)
                else:
                    body.append(p)

            except Exception as e:
                print(f"还原公式时出错: {e}")
                traceback.print_exc()
                paragraph = self.doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.add_run(f"[公式] {label}")
        else:
            paragraph = self.doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run(f"[公式] {label}")
    
    def _restore_formula_inline_multi(self, element: Dict[str, Any]):
        segments = element.get('formula_segments', [])
        if not segments:
            return
        try:
            paragraph = self.doc.add_paragraph()
            self._apply_paragraph_style(paragraph, element.get('style', {}))
            for i, seg in enumerate(segments):
                text_before = seg.get('text_before', '')
                omath_str = seg.get('omath', '')
                text_after = seg.get('text_after', '')
                if text_before:
                    paragraph.add_run(text_before)
                if omath_str:
                    try:
                        omath_elem = etree.fromstring(omath_str.encode())
                        paragraph._p.append(omath_elem)
                    except Exception as e:
                        paragraph.add_run('[公式]')
                if i == len(segments) - 1 and text_after:
                    paragraph.add_run(text_after)
        except Exception as e:
            print(f"还原多行内公式时出错: {e}")
            traceback.print_exc()
            paragraph = self.doc.add_paragraph()
            paragraph.add_run('[多行内公式还原失败]')

    def _restore_ole_formula(self, element: Dict[str, Any]):
        ole_base64 = element.get("ole_base64", "")
        image_base64 = element.get("image_base64", "")
        label = element.get("label", "")
        prog_id = element.get("prog_id", "Equation.3")
        width_pt = element.get("width_pt", 50)
        height_pt = element.get("height_pt", 20)
        
        if not ole_base64:
            paragraph = self.doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run(f"[公式占位符] {label}")
            return
        
        try:
            ole_bytes = base64.b64decode(ole_base64)
            image_bytes = base64.b64decode(image_base64) if image_base64 else None
            
            self._ole_counter += 1
            self._shape_counter += 1
            
            ole_rid, image_rid = self._add_ole_parts(ole_bytes, image_bytes)
            
            shape_id = f"_x0000_i{1025 + self._shape_counter - 1}"
            object_id = f"_{int(1468075725 + self._ole_counter)}"
            
            p = self._create_ole_paragraph(ole_rid, image_rid, shape_id, object_id, 
                                           prog_id, width_pt, height_pt, label)
            
            body = self.doc.element.body
            sectPr = body.find(f"{{{W}}}sectPr")
            if sectPr is not None:
                sectPr.addprevious(p)
            else:
                body.append(p)
                
        except Exception as e:
            print(f"还原OLE公式时出错: {e}")
            traceback.print_exc()
            paragraph = self.doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run(f"[公式还原失败: {e}] {label}")
    
    def _add_ole_parts(self, ole_bytes: bytes, image_bytes: bytes = None) -> tuple:
        ole_rid = self._add_ole_embedded_part(ole_bytes)
        image_rid = self._add_ole_image_part(image_bytes)
        return ole_rid, image_rid
    
    def _get_next_rid(self) -> str:
        rels = self.doc.part.rels
        rids = [r.rId for r in rels.values()]
        return f"rId{max([int(r[3:]) for r in rids if r.startswith('rId')], default=0) + 1}"
    
    def _add_ole_embedded_part(self, ole_bytes: bytes) -> str:
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        
        partname = PackURI(f"/word/embeddings/oleObject{self._ole_counter}.bin")
        content_type = "application/vnd.ms-office.activeX"
        
        part = Part(partname, content_type, ole_bytes)
        
        new_rid = self._get_next_rid()
        rel = self.doc.part.rels.add_relationship(RT.OLE_OBJECT, part, new_rid)
        
        return new_rid
    
    def _add_ole_image_part(self, image_bytes: bytes = None) -> str:
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        
        if image_bytes is None:
            image_bytes = self._create_ole_placeholder_image()
        
        partname = PackURI(f"/word/media/oleImage{self._ole_counter}.wmf")
        content_type = "image/x-wmf"
        
        part = Part(partname, content_type, image_bytes)
        
        new_rid = self._get_next_rid()
        rel = self.doc.part.rels.add_relationship(RT.IMAGE, part, new_rid)
        
        return new_rid
    
    def _create_ole_placeholder_image(self) -> bytes:
        wmf_header = bytes([
            0xD7, 0xCD, 0xC6, 0x9A,
            0x00, 0x00,
            0x00, 0x00, 0x00, 0x00,
            0x00, 0x00,
            0x00, 0x00,
            0x00, 0x00, 0x00, 0x00,
            0x00, 0x00, 0x00, 0x00,
        ])
        return wmf_header
    
    def _create_ole_paragraph(self, ole_rid: str, image_rid: str, shape_id: str, 
                              object_id: str, prog_id: str, width_pt: float, 
                              height_pt: float, label: str) -> etree._Element:
        p = etree.Element(f"{{{W}}}p")
        
        pPr = etree.SubElement(p, f"{{{W}}}pPr")
        
        jc = etree.SubElement(pPr, f"{{{W}}}jc")
        jc.set(f"{{{W}}}val", "center")
        
        r = etree.SubElement(p, f"{{{W}}}r")
        
        obj = etree.SubElement(r, f"{{{W}}}object")
        
        shape = etree.SubElement(obj, f"{{{V_NS}}}shape")
        shape.set("id", shape_id)
        shape.set(f"{{{O_NS}}}spt", "75")
        shape.set("type", "#_x0000_t75")
        shape.set("style", f"height:{height_pt}pt;width:{width_pt}pt;")
        shape.set(f"{{{O_NS}}}ole", "t")
        shape.set("filled", "f")
        shape.set(f"{{{O_NS}}}preferrelative", "t")
        shape.set("stroked", "f")
        shape.set("coordsize", "21600,21600")
        
        etree.SubElement(shape, f"{{{V_NS}}}path")
        
        fill = etree.SubElement(shape, f"{{{V_NS}}}fill")
        fill.set("on", "f")
        fill.set("focussize", "0,0")
        
        stroke = etree.SubElement(shape, f"{{{V_NS}}}stroke")
        stroke.set("on", "f")
        
        imagedata = etree.SubElement(shape, f"{{{V_NS}}}imagedata")
        imagedata.set(f"{{{R_NS}}}id", image_rid)
        imagedata.set(f"{{{O_NS}}}title", "")
        
        lock = etree.SubElement(shape, f"{{{O_NS}}}lock")
        lock.set(f"{{{V_NS}}}ext", "edit")
        lock.set("grouping", "f")
        lock.set("rotation", "f")
        lock.set("text", "f")
        lock.set("aspectratio", "t")
        
        wrap = etree.SubElement(shape, f"{{{W}}}wrap")
        wrap.set("type", "none")
        
        anchorlock = etree.SubElement(shape, f"{{{W}}}anchorlock")
        
        ole_obj = etree.SubElement(obj, f"{{{O_NS}}}OLEObject")
        ole_obj.set("Type", "Embed")
        ole_obj.set("ProgID", prog_id)
        ole_obj.set("ShapeID", shape_id)
        ole_obj.set("DrawAspect", "Content")
        ole_obj.set("ObjectID", object_id)
        ole_obj.set(f"{{{R_NS}}}id", ole_rid)
        
        locked_field = etree.SubElement(ole_obj, f"{{{O_NS}}}LockedField")
        locked_field.text = "false"
        
        if label:
            r_label = etree.SubElement(p, f"{{{W}}}r")
            t_label = etree.SubElement(r_label, f"{{{W}}}t")
            t_label.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t_label.text = f"    {label}"
        
        return p
    
    def restore(self) -> Document:
        for element in self.elements:
            elem_type = element.get("type", "body")
            
            if elem_type.startswith("heading"):
                self._restore_paragraph(element)
            elif elem_type == "body":
                self._restore_paragraph(element)
            elif elem_type == "table":
                self._restore_table(element)
            elif elem_type == "image":
                self._restore_image(element)
            elif elem_type in ("formula", "formula_block", "formula_inline"):
                self._restore_formula(element)
            elif elem_type == "formula_inline_multi":
                self._restore_formula_inline_multi(element)
            else:
                self._restore_paragraph(element)
        
        if self.citations:
            self._restore_citations()
        
        return self.doc
    
    def save(self, output_path: str):
        self.restore()
        self.doc.save(output_path)


def restore_docx(elements: List[Dict[str, Any]], output_docx_path: str, citations: List[Dict[str, Any]] = None) -> Document:
    """
    从JSON元素列表还原docx文档
    
    Args:
        elements: JSON元素列表
        output_docx_path: 输出docx文件路径
        citations: 引用配置列表（可选）
    
    Returns:
        还原后的Document对象
    """
    restorer = DocxRestorer(elements, citations)
    restorer.save(output_docx_path)
    return restorer.doc


def restore_from_json(json_path: str, output_docx_path: str, citations_path: str = None) -> Document:
    """
    从JSON文件还原docx文档
    
    Args:
        json_path: JSON文件路径
        output_docx_path: 输出docx文件路径
        citations_path: 引用配置文件路径（可选）
    
    Returns:
        还原后的Document对象
    """
    with open(json_path, "r", encoding="utf-8") as f:
        elements = json.load(f)
    
    citations = None
    if citations_path and Path(citations_path).exists():
        with open(citations_path, "r", encoding="utf-8") as f:
            citations = json.load(f)
    
    return restore_docx(elements, output_docx_path, citations)


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 3:
        print("用法: python docx_restorer.py <JSON文件路径> <输出docx路径>")
        sys.exit(1)
    
    json_path = sys.argv[1]
    output_path = sys.argv[2]
    
    try:
        restore_from_json(json_path, output_path)
        print(f"还原完成，文档已保存到: {output_path}")
    except Exception as e:
        print(f"还原失败: {e}")
        traceback.print_exc()
